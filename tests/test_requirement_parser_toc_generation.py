from __future__ import annotations

import unittest

from docx import Document

from backend.agents.requirement_parser import RequirementParserAgent
from backend.agents.toc_generator import TOCGeneratorAgent
from backend.models.schemas import TOCNode
from tests.helpers import cleanup_temp_root, create_sample_docx, make_temp_root


class RequirementParserAndTOCTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_root = make_temp_root("req_toc_test_")

    def tearDown(self) -> None:
        cleanup_temp_root(self.temp_root)

    def test_parser_is_clause_aware_and_extracts_subsystems(self) -> None:
        docx = create_sample_docx(
            self.temp_root / "sample.docx",
            [
                "第五章 服务要求",
                "3.13.1.1 优化和完善太阳能余热回收控制系统，根据生产管理的要求，对太阳能余热综合利用系统进行开发和优化，包括太阳能供热子系统、空气源热泵子系统。",
                "3.13.1.2 优化和完善太阳能余热回收控制系统，当前的太阳能余热回收控制系统由一期（卷包）和二期（制丝）组成，要求优化现有控制逻辑并完成接口联调。",
                "验收应符合GB50348。",
            ],
        )

        requirement, parse_report = RequirementParserAgent().parse(
            task_id="task_test",
            upload_file_path=docx,
            fallback_title="太阳能余热综合利用项目",
        )

        subsystem_names = [item.name for item in requirement.scope.subsystems]
        self.assertEqual(requirement.project.name, "太阳能余热回收控制系统")
        self.assertIn("太阳能余热回收控制系统", subsystem_names)
        self.assertIn("太阳能供热子系统", subsystem_names)
        self.assertIn("空气源热泵子系统", subsystem_names)
        self.assertNotIn("优化和完善太阳能余热回收控制系统", subsystem_names)
        self.assertIn("GB50348", requirement.constraints.standards)
        self.assertEqual(parse_report["subsystem_count"], 3)

    def test_toc_generator_prefers_doc_heading_hierarchy(self) -> None:
        docx = self.temp_root / "outline.docx"
        document = Document()
        document.add_paragraph("空压站余热综合利用实施方案", style="Heading 1")
        document.add_paragraph("建设内容", style="Heading 2")
        document.add_paragraph("空压机余热回收系统", style="Heading 3")
        document.add_paragraph("现场实施条件", style="Heading 4")
        document.add_paragraph("应完成现场勘察、基础复核和施工边界确认。")
        document.add_paragraph("控制系统联调", style="Heading 3")
        document.add_paragraph("需与既有 PLC 和上位系统完成接口联调。")
        document.save(docx)

        requirement, _ = RequirementParserAgent().parse(
            task_id="task_outline",
            upload_file_path=docx,
            fallback_title="空压站余热综合利用实施方案",
        )
        toc = TOCGeneratorAgent().generate(requirement=requirement, version_no=1)

        root = toc.tree[0]
        self.assertEqual(root.title, "空压站余热综合利用实施方案")
        self.assertEqual([child.title for child in root.children], ["建设内容"])
        self.assertEqual([child.title for child in root.children[0].children], ["空压机余热回收系统", "控制系统联调"])
        self.assertEqual(root.children[0].children[0].children[0].title, "现场实施条件")
        self.assertTrue(root.children[0].children[0].children[0].is_generation_unit)

    def test_toc_generator_recognizes_numbered_heading_hierarchy_without_styles(self) -> None:
        docx = create_sample_docx(
            self.temp_root / "outline_numbered.docx",
            [
                "1 空压站余热综合利用实施方案",
                "1.1 建设内容",
                "1.1.1 空压机余热回收系统",
                "1.1.1.1 现场实施条件",
                "1.1.1.2 控制系统联调",
                "3.13.1.1 优化和完善太阳能余热回收控制系统，根据生产管理的要求，对太阳能余热综合利用系统进行开发和优化。",
            ],
        )

        requirement, _ = RequirementParserAgent().parse(
            task_id="task_outline_numbered",
            upload_file_path=docx,
            fallback_title="空压站余热综合利用实施方案",
        )
        toc = TOCGeneratorAgent().generate(requirement=requirement, version_no=1)

        root = toc.tree[0]
        self.assertEqual(root.title, "空压站余热综合利用实施方案")
        self.assertEqual(root.children[0].title, "建设内容")
        self.assertEqual(root.children[0].children[0].title, "空压机余热回收系统")
        self.assertEqual(
            [child.title for child in root.children[0].children[0].children],
            ["现场实施条件", "控制系统联调"],
        )

    def test_outline_leaf_sections_use_topic_specific_children_instead_of_generic_phases(self) -> None:
        generator = TOCGeneratorAgent()
        section_map = {
            "验收标准": generator._build_outline_leaf_specs("验收标准"),
            "水源热泵技术要求": generator._build_outline_leaf_specs("水源热泵技术要求"),
            "设备及管路系统安装要求": generator._build_outline_leaf_specs("设备及管路系统安装要求"),
            "其他": generator._build_outline_leaf_specs("其他"),
        }
        self.assertEqual(
            section_map["验收标准"],
            ["验收依据与判定标准", "功能与性能验收要求", "验收资料与签认要求"],
        )
        self.assertEqual(
            section_map["水源热泵技术要求"],
            ["水源热泵配置与选型要求", "水源热泵性能与控制要求", "水源热泵安装与验收要求"],
        )
        self.assertEqual(
            section_map["设备及管路系统安装要求"],
            ["设备及管路系统安装范围与条件", "设备及管路系统安装工艺与质量控制", "设备及管路系统调试与验收要求"],
        )
        self.assertEqual(
            section_map["其他"],
            ["补充实施要求", "配合条件与边界控制", "资料提交与后续支持要求"],
        )

    def test_generator_builds_comprehensive_engineering_outline_for_energy_retrofit_doc(self) -> None:
        docx = create_sample_docx(
            self.temp_root / "engineering_outline.docx",
            [
                "第五章 服务要求",
                "一、项目内容",
                "在大理卷烟厂厂区内新建1套空压机余热回收系统、1套乏汽余热回收系统、1套水箱换热系统，更新2组空气源热泵、1台管壳式换热器，并对空压机群控系统、太阳能余热回收控制系统及蓄热器排污余热回收系统进行优化。",
                "购置并安装一台冷侧制冷量不低于750KW的水源热泵机组。",
                "购置并安装14台单台制热量不低于70kW的空气源热泵。",
                "购置并安装一个容积不低于28m³的热分层储能装置。",
                "购置并安装2台换热量不低于800kW的板式换热器。",
                "购置并安装1台换热量不低于1100kW的闪蒸汽（乏汽）回收装置。",
                "购置并安装1台换热量不低于600kW的管式换热器，安装在制丝车间热力站内。",
                "完成空压机余热回收系统、太阳能系统在上位系统集成。",
                "二、相关标准规范",
                "2.1 设计安装标准",
                "2.2 验收标准",
                "三、项目系统技术要求",
                "3.3 水源热泵技术要求",
                "3.4 空气源热泵技术要求",
                "3.5 换热器技术要求",
                "3.8 管道技术要求",
                "3.9 阀门技术要求",
                "3.10 支架及支座技术要求",
                "3.11 保温技术要求",
                "3.12 设备及管路系统安装要求",
                "3.13 电控系统技术要求",
                "3.14 上位系统集成要求",
                "3.15 其他",
                "四、项目主要材料清单",
            ],
        )

        requirement, _ = RequirementParserAgent().parse(
            task_id="task_engineering_outline",
            upload_file_path=docx,
            fallback_title="大理卷烟厂余热综合利用项目",
        )
        toc = TOCGeneratorAgent().generate(requirement=requirement, version_no=1)

        root = toc.tree[0]
        chapter_titles = [node.title for node in root.children]
        self.assertEqual(root.title, "工程实施方案")
        self.assertEqual(chapter_titles[0], "项目理解与建设目标")
        self.assertEqual(chapter_titles[1], "项目建设内容总述")
        self.assertIn("总体技术方案", chapter_titles)
        self.assertIn("空压机余热回收系统方案", chapter_titles)
        self.assertIn("电气与自控系统方案", chapter_titles)
        self.assertGreaterEqual(len(chapter_titles), 15)

        chapter5 = next(node for node in root.children if node.title == "总体技术方案")
        self.assertEqual(
            [child.title for child in chapter5.children[:3]],
            ["系统总体架构", "热源综合利用逻辑", "各子系统耦合关系"],
        )
        self.assertEqual(
            [child.title for child in chapter5.children[0].children],
            ["热源侧系统架构", "用户侧系统架构", "自控与监控系统架构"],
        )

    def test_generator_does_not_force_same_full_template_for_small_doc(self) -> None:
        docx = create_sample_docx(
            self.temp_root / "small_engineering_outline.docx",
            [
                "项目内容",
                "新建一套视频监控系统，完成前端摄像机、存储和平台接入。",
                "相关标准规范",
                "设计安装标准",
                "系统技术要求",
                "视频监控系统需支持联网接入和远程查看。",
                "验收标准",
            ],
        )

        requirement, _ = RequirementParserAgent().parse(
            task_id="task_small_outline",
            upload_file_path=docx,
            fallback_title="视频监控系统项目",
        )
        toc = TOCGeneratorAgent().generate(requirement=requirement, version_no=1)

        chapter_titles = [node.title for node in toc.tree[0].children]
        self.assertGreaterEqual(len(chapter_titles), 5)
        self.assertLess(len(chapter_titles), 20)
        self.assertIn("项目理解与建设目标", chapter_titles)
        self.assertIn("设计依据与执行标准", chapter_titles)
        self.assertIn("调试、试运行与验收方案", chapter_titles)
        self.assertNotIn("乏汽余热回收系统方案", chapter_titles)
        self.assertNotIn("空气源热泵系统方案", chapter_titles)

    def test_toc_generator_builds_engineering_tree_with_stable_node_uid(self) -> None:
        docx = create_sample_docx(
            self.temp_root / "sample.docx",
            [
                "第五章 服务要求",
                "3.13.1.1 优化和完善太阳能余热回收控制系统，根据生产管理的要求，对太阳能余热综合利用系统进行开发和优化，包括太阳能供热子系统、空气源热泵子系统。",
                "3.13.1.2 优化和完善太阳能余热回收控制系统，当前的太阳能余热回收控制系统由一期（卷包）和二期（制丝）组成，要求优化现有控制逻辑并完成接口联调。",
                "验收应符合GB50348。",
            ],
        )
        requirement, _ = RequirementParserAgent().parse(
            task_id="task_test",
            upload_file_path=docx,
            fallback_title="太阳能余热综合利用项目",
        )

        generator = TOCGeneratorAgent()
        toc_v1 = generator.generate(requirement=requirement, version_no=1)
        toc_v2 = generator.generate(requirement=requirement, version_no=2)

        root_v1 = toc_v1.tree[0]
        level2_titles = [node.title for node in root_v1.children]
        self.assertEqual(root_v1.title, "工程实施方案")
        self.assertGreaterEqual(len(root_v1.children), 3)
        self.assertIn("太阳能余热回收控制系统", level2_titles)
        self.assertIn("太阳能供热子系统", level2_titles)
        self.assertIn("空气源热泵子系统", level2_titles)

        first_level2 = root_v1.children[0]
        self.assertEqual(
            [child.title for child in first_level2.children],
            [
                "实施范围与目标",
                "实施准备与条件确认",
                "太阳能余热回收控制系统优化与接口联调",
                "联调测试与验收",
            ],
        )
        self.assertTrue(all(len(child.children) == 2 for child in first_level2.children))
        self.assertTrue(all(grandchild.is_generation_unit for child in first_level2.children for grandchild in child.children))

        self.assertEqual(
            self._flatten_node_uids(toc_v1.tree),
            self._flatten_node_uids(toc_v2.tree),
        )

    def _flatten_node_uids(self, nodes: list[TOCNode]) -> list[str]:
        values: list[str] = []
        for node in nodes:
            values.append(node.node_uid)
            values.extend(self._flatten_node_uids(node.children))
        return values


if __name__ == "__main__":
    unittest.main()
