from __future__ import annotations

import json
import shutil
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from PIL import Image

from backend.agents.layout import LayoutAgent
from backend.agents.word_export import WordExportAgent
from backend.models.schemas import TOCDocument, TOCNode


class LayoutExportTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_root = Path(tempfile.mkdtemp(prefix="layout_export_test_", dir="/tmp"))
        self.artifacts_root = self.temp_root / "artifacts"
        self.task_id = "task_test"
        self.node_dir = self.artifacts_root / self.task_id / "nodes" / "uid_001"
        (self.node_dir / "images").mkdir(parents=True, exist_ok=True)
        (self.artifacts_root / self.task_id / "toc").mkdir(parents=True, exist_ok=True)
        (self.artifacts_root / self.task_id / "final").mkdir(parents=True, exist_ok=True)
        Image.new("RGB", (900, 1600), color=(180, 200, 230)).save(
            self.node_dir / "images" / "img_001.png"
        )

        (self.node_dir / "text.json").write_text(
            json.dumps(
                {
                    "node_uid": "uid_001",
                    "node_id": "1.1.1",
                    "title": "循环水泵配置方案",
                    "summary": "用于验证排版与导出。",
                    "sections": [
                        {
                            "section_id": "sec_01",
                            "title": "实施步骤",
                            "paragraphs": [
                                {
                                    "paragraph_id": "p_01",
                            "text": "本节说明循环水泵配置、安装要点、质量控制与验收留痕要求。",
                                    "source_refs": [],
                                    "claim_ids": [],
                                    "anchors": ["anchor_impl"],
                                }
                            ],
                        }
                    ],
                    "highlight_paragraphs": [],
                    "word_count": 24,
                    "version": 1,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        (self.node_dir / "tables.json").write_text(
            json.dumps(
                {
                    "node_uid": "uid_001",
                    "tables": [
                        {
                            "table_id": "table_01",
                            "title": "主要设备配置表",
                            "headers": ["设备", "数量"],
                            "rows": [["交换机", "2"], ["机柜", "1"]],
                            "style_name": "BiddingTable",
                            "bind_anchor": "anchor_impl",
                            "source_refs": [],
                        }
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        (self.node_dir / "images.json").write_text(
            json.dumps(
                {
                    "node_uid": "uid_001",
                    "images": [
                        {
                            "image_id": "img_001",
                            "type": "process",
                            "file": "images/img_001.png",
                            "caption": "图001 循环水泵示意",
                            "style_preset": "engineering_simulation_detail",
                            "aspect_ratio": "2:1",
                            "group_caption": None,
                            "bind_anchor": "anchor_impl",
                            "bind_section": "实施步骤",
                            "status": "PASS",
                            "retry_count": 0,
                        }
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

    def tearDown(self) -> None:
        shutil.rmtree(self.temp_root)

    def _build_toc_document(self) -> TOCDocument:
        return TOCDocument(
            version=1,
            tree=[
                TOCNode(
                    node_uid="uid_root",
                    node_id="1",
                    level=1,
                    title="工程实施方案",
                    is_generation_unit=False,
                    source_refs=[],
                    constraints=None,
                    children=[
                        TOCNode(
                            node_uid="uid_l2",
                            node_id="1.1",
                            level=2,
                            title="施工组织与实施方案",
                            is_generation_unit=False,
                            source_refs=[],
                            constraints=None,
                            children=[
                                TOCNode(
                                    node_uid="uid_l3",
                                    node_id="1.1.1",
                                    level=3,
                                    title="设备安装方案",
                                    is_generation_unit=False,
                                    source_refs=[],
                                    constraints=None,
                                    children=[
                                        TOCNode(
                                            node_uid="uid_001",
                                            node_id="1.1.1.1",
                                            level=4,
                                            title="循环水泵配置方案",
                                            is_generation_unit=True,
                                            source_refs=[],
                                            constraints=None,
                                            children=[],
                                        )
                                    ],
                                )
                            ],
                        )
                    ],
                )
            ],
        )

    def test_layout_and_export_generate_docx(self) -> None:
        toc_document = self._build_toc_document()

        layout_agent = LayoutAgent()
        payload = layout_agent.build(
            task_id=self.task_id,
            artifacts_root=self.artifacts_root,
            toc_document=toc_document,
        )
        layout_path = self.artifacts_root / self.task_id / "final" / "layout_blocks.json"
        layout_path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        export_agent = WordExportAgent()
        output_path = self.artifacts_root / self.task_id / "final" / "output.docx"
        export_agent.export(
            template_path=Path("templates/standard_template.docx"),
            layout_blocks_path=layout_path,
            output_path=output_path,
        )

        self.assertTrue(layout_path.exists())
        self.assertTrue(output_path.exists())

        doc = Document(output_path)
        self.assertGreaterEqual(len(doc.paragraphs), 5)
        self.assertEqual(len(doc.tables), 1)
        self.assertEqual(len(doc.inline_shapes), 1)
        self.assertEqual(doc.paragraphs[0].text, "施工组织与实施方案")
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")
        self.assertEqual(doc.paragraphs[1].text, "设备安装方案")
        self.assertEqual(doc.paragraphs[1].style.name, "Heading 2")
        self.assertEqual(doc.paragraphs[2].text, "循环水泵配置方案")
        self.assertEqual(doc.paragraphs[2].style.name, "Heading 3")
        self.assertEqual(doc.paragraphs[3].text, "实施步骤")
        self.assertEqual(doc.paragraphs[3].style.name, "Normal")
        self.assertEqual(doc.paragraphs[3].paragraph_format.line_spacing, 1.5)
        self.assertEqual(doc.paragraphs[3].runs[0].font.color.rgb, RGBColor(0x00, 0x70, 0xC0))
        self.assertEqual(doc.paragraphs[4].style.name, "Normal")
        self.assertIsNotNone(doc.paragraphs[4].paragraph_format.first_line_indent)
        self.assertAlmostEqual(
            doc.paragraphs[4].paragraph_format.first_line_indent.pt,
            24.0,
            places=1,
        )
        self.assertEqual(doc.paragraphs[4].paragraph_format.line_spacing, 1.5)
        self.assertEqual(doc.tables[0].style.name, "BiddingTable")
        self.assertTrue(all("图001" not in para.text for para in doc.paragraphs))
        self.assertNotIn("<w:numPr>", doc.paragraphs[3]._p.xml)
        self.assertNotIn("<w:numPr>", doc.paragraphs[4]._p.xml)
        image_paragraph = next((para for para in doc.paragraphs if not para.text.strip()), None)
        self.assertIsNotNone(image_paragraph)
        self.assertNotIn("<w:numPr>", image_paragraph._p.xml)
        section = doc.sections[-1]
        available_width = int(section.page_width - section.left_margin - section.right_margin)
        image_width = int(doc.inline_shapes[0].width)
        image_height = int(doc.inline_shapes[0].height)
        self.assertGreaterEqual(image_width, int(available_width * 0.99))
        self.assertLessEqual(image_width, available_width)
        self.assertAlmostEqual(image_width / image_height, 2.0, delta=0.03)
        self.assertEqual(doc.tables[0].cell(0, 0).vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
        self.assertEqual(doc.tables[0].cell(0, 0).paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(doc.tables[0].cell(1, 1).paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_layout_and_export_generate_docx_without_images(self) -> None:
        layout_agent = LayoutAgent()
        payload = layout_agent.build(
            task_id=self.task_id,
            artifacts_root=self.artifacts_root,
            toc_document=self._build_toc_document(),
            include_images=False,
        )
        layout_path = self.artifacts_root / self.task_id / "final" / "layout_blocks.json"
        layout_path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        export_agent = WordExportAgent()
        output_path = self.artifacts_root / self.task_id / "final" / "output_no_images.docx"
        export_agent.export(
            template_path=Path("templates/standard_template.docx"),
            layout_blocks_path=layout_path,
            output_path=output_path,
        )

        doc = Document(output_path)
        self.assertEqual(len(doc.inline_shapes), 0)


if __name__ == "__main__":
    unittest.main()
