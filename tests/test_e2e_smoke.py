from __future__ import annotations

import json
import unittest
from unittest.mock import patch

from docx import Document

from backend.agents.requirement_parser import RequirementParserAgent
from backend.agents.section_writer import SectionWriterAgent
from backend.agents.toc_generator import TOCGeneratorAgent
from backend.agents.toc_review.agent import TOCReviewChatAgent
from backend.app_service.task_service import TaskService
from backend.models.enums import NodeStatus, TaskStatus
from backend.worker.worker_process import WorkerProcess
from tests.helpers import build_settings, cleanup_temp_root, create_sample_docx, make_temp_root


class EndToEndSmokeTestCase(unittest.TestCase):
    _MODEL_TOC = """
    {
      "root_title": "视频监控系统实施方案",
      "chapters": [
        {
          "title": "视频监控系统建设范围",
          "children": [
            {
              "title": "前端建设要求",
              "children": [
                {"title": "前端点位部署要求", "children": []}
              ]
            },
            {
              "title": "平台接入要求",
              "children": [
                {"title": "平台接入与联调要求", "children": []}
              ]
            }
          ]
        },
        {
          "title": "施工与验收要求",
          "children": [
            {
              "title": "施工管理要求",
              "children": [
                {"title": "施工组织与质量控制", "children": []}
              ]
            },
            {
              "title": "验收交付要求",
              "children": [
                {"title": "验收资料与签认要求", "children": []}
              ]
            }
          ]
        }
      ]
    }
    """

    @staticmethod
    def _review_model_actions(target_title: str, renamed_title: str) -> str:
        return json.dumps(
            {
                "actions": [
                    {
                        "type": "rename",
                        "target": target_title,
                        "new_title": renamed_title,
                    }
                ]
            },
            ensure_ascii=False,
        )

    @staticmethod
    def _repeat_to_length(text: str, target_words: int) -> str:
        repeat = max(1, (target_words // max(1, len(text))) + 2)
        return (text * repeat)[:target_words]

    @classmethod
    def _video_section_writer_output(cls) -> str:
        paragraph_a = cls._repeat_to_length(
            "视频监控子系统应完成前端部署、链路联调和平台接入，并以招标文件要求组织实施和复核。",
            980,
        )
        paragraph_b = cls._repeat_to_length(
            "施工过程应符合GB50348，验收阶段应形成记录并完成签认，相关资料应按原文要求整理归档。",
            980,
        )
        return json.dumps(
            {
                "summary": "根据招标文件全文提炼当前节点实施要求。",
                "sections": [
                    {
                        "title": "实施要求与控制要点",
                        "paragraphs": [
                            {
                                "text": paragraph_a,
                                "source_refs": ["p1#L1"],
                            },
                            {
                                "text": paragraph_b,
                                "source_refs": ["p1#L2", "p1#L3"],
                            },
                        ],
                    }
                ],
                "highlight_paragraphs": ["关键过程和验收资料必须来源可追溯。"],
            },
            ensure_ascii=False,
        )

    @classmethod
    def _service_section_writer_output(cls) -> str:
        paragraph_a = cls._repeat_to_length(
            "售后服务需覆盖维保、巡检、培训和应急响应，所有动作均应依据原文范围组织实施和留痕。",
            980,
        )
        paragraph_b = cls._repeat_to_length(
            "实施过程应兼顾不停机生产与分阶段改造窗口，应急处置需具备远程诊断、现场联动和闭环整改能力。",
            980,
        )
        return json.dumps(
            {
                "summary": "根据招标文件全文提炼当前节点服务要求。",
                "sections": [
                    {
                        "title": "服务要求与响应边界",
                        "paragraphs": [
                            {
                                "text": paragraph_a,
                                "source_refs": ["p1#L1"],
                            },
                            {
                                "text": paragraph_b,
                                "source_refs": ["p1#L2", "p1#L3"],
                            },
                        ],
                    }
                ],
                "highlight_paragraphs": ["服务要求必须严格来自招标文件原文。"],
            },
            ensure_ascii=False,
        )

    @staticmethod
    def _parser_model_response(*args, **kwargs) -> str:
        prompt = kwargs["prompt"]
        if "整编 requirement.json 的核心摘要字段" in prompt and "视频监控子系统" in prompt:
            return json.dumps(
                {
                    "project_name": "端到端烟测项目",
                    "overview": "视频监控子系统需完成前端部署、链路联调、平台接入，并满足施工与验收要求。",
                    "subsystems": [
                        {
                            "name": "视频监控子系统",
                            "description": "完成前端部署、链路联调和平台接入。",
                            "source_refs": ["p1#L2"],
                        }
                    ],
                    "standards": ["GB50348"],
                    "acceptance": ["验收阶段应形成记录并完成签认。"],
                },
                ensure_ascii=False,
            )
        if "整编 requirement.json 的核心摘要字段" in prompt:
            return json.dumps(
                {
                    "project_name": "手工目录树烟测项目",
                    "overview": "售后服务需覆盖维保、巡检、培训和应急响应，并兼顾不停机生产与分阶段改造。",
                    "subsystems": [
                        {
                            "name": "售后服务",
                            "description": "覆盖维保、巡检、培训和应急响应。",
                            "source_refs": ["p1#L2", "p1#L3", "p1#L4"],
                        }
                    ],
                    "standards": [],
                    "acceptance": [],
                },
                ensure_ascii=False,
            )
        if "视频监控子系统" in prompt:
            return json.dumps(
                {
                    "overview_points": ["视频监控子系统实施与验收要求。"],
                    "requirements": [
                        {
                            "type": "technical",
                            "key": "video_deploy",
                            "value": "视频监控子系统需要完成前端部署、链路联调和平台接入。",
                            "source_ref": "p1#L2",
                        },
                        {
                            "type": "standard",
                            "key": "gb50348",
                            "value": "GB50348",
                            "source_ref": "p1#L3",
                        },
                        {
                            "type": "acceptance",
                            "key": "acceptance_record",
                            "value": "验收阶段应形成记录并完成签认。",
                            "source_ref": "p1#L4",
                        },
                    ],
                    "subsystems": [
                        {
                            "name": "视频监控子系统",
                            "description": "完成前端部署、链路联调和平台接入。",
                            "source_refs": ["p1#L2"],
                        }
                    ],
                    "standards": [{"name": "GB50348", "source_ref": "p1#L3"}],
                    "acceptance": [{"value": "验收阶段应形成记录并完成签认。", "source_ref": "p1#L4"}],
                },
                ensure_ascii=False,
            )
        return json.dumps(
            {
                "overview_points": ["售后服务与应急响应要求。"],
                "requirements": [
                    {
                        "type": "service",
                        "key": "service_scope",
                        "value": "售后服务需覆盖维保、巡检、培训和应急响应。",
                        "source_ref": "p1#L2",
                    },
                    {
                        "type": "operation",
                        "key": "window_constraint",
                        "value": "实施过程应兼顾不停机生产与分阶段改造窗口。",
                        "source_ref": "p1#L3",
                    },
                    {
                        "type": "service",
                        "key": "emergency_loop",
                        "value": "应急处置需具备远程诊断、现场联动和闭环整改能力。",
                        "source_ref": "p1#L4",
                    },
                ],
                "subsystems": [
                    {
                        "name": "售后服务",
                        "description": "覆盖维保、巡检、培训和应急响应。",
                        "source_refs": ["p1#L2", "p1#L3", "p1#L4"],
                    }
                ],
                "standards": [],
                "acceptance": [],
            },
            ensure_ascii=False,
        )

    @staticmethod
    def _manual_outline_model_toc() -> str:
        return json.dumps(
            {
                "root_title": "售后服务实施方案",
                "chapters": [
                    {
                        "title": "售后服务总体方案",
                        "children": [
                            {
                                "title": "售后服务目标",
                                "children": [{"title": "保障系统安全稳定运行", "children": []}],
                            }
                        ],
                    },
                    {
                        "title": "应急响应措施",
                        "children": [
                            {
                                "title": "应急响应总体机制",
                                "children": [{"title": "全天候响应机制", "children": []}],
                            }
                        ],
                    },
                ],
            },
            ensure_ascii=False,
        )

    def setUp(self) -> None:
        self.temp_root = make_temp_root("e2e_smoke_test_")
        self.settings = build_settings(self.temp_root)
        self.service = TaskService(settings=self.settings)
        self.worker = WorkerProcess(self.settings)
        self.service.update_system_config(
            {
                "text_provider": "minimax",
                "text_model_name": "MiniMax-M2.5",
                "text_api_key": "fake-key",
            }
        )

    def tearDown(self) -> None:
        cleanup_temp_root(self.temp_root)

    def test_docx_to_output_smoke_flow(self) -> None:
        task = self.service.create_task("端到端烟测")
        source_docx = create_sample_docx(
            self.temp_root / "input" / "sample.docx",
            [
                "端到端烟测项目",
                "视频监控子系统需要完成前端部署、链路联调和平台接入。",
                "施工过程应符合GB50348。",
                "验收阶段应形成记录并完成签认。",
            ],
        )

        with (
            patch.object(
                RequirementParserAgent,
                "_request_minimax_completion",
                side_effect=self._parser_model_response,
            ),
            patch.object(TOCGeneratorAgent, "_request_minimax_completion", return_value=self._MODEL_TOC),
            patch.object(
                TOCReviewChatAgent,
                "_request_minimax_completion",
                return_value=self._review_model_actions("视频监控系统建设范围", "视频监控系统建设范围（修订版）"),
            ),
            patch.object(
                SectionWriterAgent,
                "_request_minimax_completion",
                return_value=self._video_section_writer_output(),
            ),
        ):
            upload_path = self.service.save_upload(task.task_id, "sample.docx", source_docx.read_bytes())
            parse_payload = self.service.parse_requirement(task.task_id)
            v1 = self.service.generate_toc(task.task_id)
            toc_v1 = self.service.get_toc_document(task.task_id, v1.version_no)
            first_title = toc_v1["tree"][0]["children"][0]["title"]
            v2 = self.service.review_toc(task.task_id, f'把“{first_title}”改成“{first_title}（修订版）”')
            confirm_payload = self.service.confirm_and_start_generation(task.task_id, v2.version_no)
            seeded_nodes = confirm_payload["seeded_nodes"]
            processed = self.worker.run_once()

        final_task = self.service.get_task(task.task_id)
        nodes = self.service.get_node_states(task.task_id)
        logs = self.service.get_event_logs(task.task_id, limit=200)
        output_path = self.service.get_output_path(task.task_id)

        self.assertTrue(upload_path.exists())
        self.assertEqual(parse_payload["task_id"], task.task_id)
        self.assertEqual(v1.version_no, 1)
        self.assertEqual(v2.version_no, 2)
        self.assertGreaterEqual(seeded_nodes, 1)
        self.assertEqual(confirm_payload["task"]["status"], TaskStatus.GENERATING.value)
        self.assertEqual(processed, 1)
        self.assertIsNotNone(final_task)
        assert final_task is not None
        self.assertEqual(final_task.status, TaskStatus.DONE)
        self.assertEqual(final_task.current_stage, "DONE")
        self.assertIsNotNone(output_path)
        assert output_path is not None
        self.assertTrue(output_path.exists())
        self.assertTrue(all(node.status == NodeStatus.NODE_DONE for node in nodes))
        self.assertTrue(any(event.stage == "DONE" for event in logs))

        exported = Document(output_path)
        self.assertGreater(len(exported.paragraphs), 0)
        self.assertGreaterEqual(len(nodes), 1)

    def test_manual_outline_to_output_smoke_flow(self) -> None:
        task = self.service.create_task("手工目录树烟测")
        source_docx = create_sample_docx(
            self.temp_root / "input" / "manual_outline.docx",
            [
                "手工目录树烟测项目",
                "售后服务需覆盖维保、巡检、培训和应急响应。",
                "实施过程应兼顾不停机生产与分阶段改造窗口。",
                "应急处置需具备远程诊断、现场联动和闭环整改能力。",
            ],
        )

        outline = "\n".join(
            [
                "一、售后服务总体方案",
                "1.1 售后服务目标",
                "1.1.1 保障系统安全稳定运行",
                "二、应急响应措施",
                "2.1 应急响应总体机制",
                "2.1.1 7×24小时响应机制",
            ]
        )

        with (
            patch.object(
                RequirementParserAgent,
                "_request_minimax_completion",
                side_effect=self._parser_model_response,
            ),
            patch.object(
                SectionWriterAgent,
                "_request_minimax_completion",
                return_value=self._service_section_writer_output(),
            ),
            patch.object(
                TOCGeneratorAgent,
                "_request_minimax_completion",
                return_value=self._manual_outline_model_toc(),
            ),
        ):
            self.service.save_upload(task.task_id, "manual_outline.docx", source_docx.read_bytes())
            imported = self.service.import_toc_outline(task.task_id, outline)
            confirm_payload = self.service.confirm_and_start_generation(task.task_id, imported.version_no)
            processed = self.worker.run_once()

        final_task = self.service.get_task(task.task_id)
        nodes = self.service.get_node_states(task.task_id)
        output_path = self.service.get_output_path(task.task_id)

        self.assertEqual(imported.version_no, 1)
        self.assertEqual(confirm_payload["task"]["status"], TaskStatus.GENERATING.value)
        self.assertEqual(processed, 1)
        self.assertIsNotNone(final_task)
        assert final_task is not None
        self.assertEqual(final_task.status, TaskStatus.DONE)
        self.assertIsNotNone(output_path)
        assert output_path is not None
        self.assertTrue(output_path.exists())
        self.assertTrue(all(node.status == NodeStatus.NODE_DONE for node in nodes))

        exported = Document(output_path)
        all_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
        self.assertIn("售后服务总体方案", all_text)
        self.assertIn("应急响应措施", all_text)


if __name__ == "__main__":
    unittest.main()
