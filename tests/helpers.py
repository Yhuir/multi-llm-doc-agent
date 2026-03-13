from __future__ import annotations

import json
import shutil
import tempfile
from pathlib import Path

from docx import Document

from backend.config import AppSettings
from backend.models.schemas import (
    RequirementConstraints,
    RequirementDocument,
    RequirementItem,
    RequirementProject,
    RequirementScope,
    RequirementSubsystem,
    SourceIndexItem,
    utc_now_iso,
)


def make_temp_root(prefix: str) -> Path:
    return Path(tempfile.mkdtemp(prefix=prefix, dir="/tmp"))


def cleanup_temp_root(path: Path) -> None:
    shutil.rmtree(path, ignore_errors=True)


def build_settings(temp_root: Path) -> AppSettings:
    return AppSettings(
        db_path=str(temp_root / "app.db"),
        artifacts_root=str(temp_root / "artifacts"),
        template_path=str(Path("templates/standard_template.docx").resolve()),
        system_config_path=str(temp_root / "artifacts" / "system_config.json"),
        api_host="127.0.0.1",
        api_port=8000,
        worker_poll_interval_sec=0.1,
    )


def create_sample_docx(path: Path, paragraphs: list[str]) -> Path:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def sample_requirement_document(project_name: str = "智慧园区项目") -> RequirementDocument:
    return RequirementDocument(
        project=RequirementProject(
            name=project_name,
            customer="示例客户",
            location="昆明",
            duration_days=90,
            milestones=[],
        ),
        scope=RequirementScope(
            overview=f"{project_name}工程实施范围说明",
            subsystems=[
                RequirementSubsystem(
                    name="视频监控子系统",
                    description="完成前端设备部署、链路联调、平台接入与验收留痕。",
                    requirements=[
                        RequirementItem(
                            type="text",
                            key="scope",
                            value="完成前端设备部署、链路联调、平台接入与验收留痕。",
                            source_ref="p1#L1",
                        )
                    ],
                    interfaces=["平台接口"],
                )
            ],
        ),
        constraints=RequirementConstraints(
            standards=["GB50348"],
            acceptance=["验收阶段应形成记录并完成签认。"],
        ),
        bidding_requirements=[
            RequirementItem(
                type="technical",
                key="front_end_deploy",
                value="完成前端设备部署、链路联调、平台接入与验收留痕。",
                source_ref="p1#L1",
            ),
            RequirementItem(
                type="standard",
                key="gb50348",
                value="GB50348",
                source_ref="p1#L2",
            ),
            RequirementItem(
                type="acceptance",
                key="acceptance_record",
                value="验收阶段应形成记录并完成签认。",
                source_ref="p1#L2",
            ),
        ],
        source_index={
            "p1#L1": SourceIndexItem(
                page=1,
                paragraph_id="para_1",
                text="视频监控子系统实施要求与验收约束。",
            ),
            "p1#L2": SourceIndexItem(
                page=1,
                paragraph_id="para_2",
                text="施工过程应符合GB50348并保留验收记录。",
            ),
        },
    )


def write_requirement_artifact(
    artifacts_root: Path,
    task_id: str,
    requirement: RequirementDocument | None = None,
) -> Path:
    target = artifacts_root / task_id / "parsed" / "requirement.json"
    payload = (requirement or sample_requirement_document()).model_dump(mode="json")
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return target


def write_style_profile_artifact(artifacts_root: Path, task_id: str) -> Path:
    target = artifacts_root / task_id / "parsed" / "style_profile.json"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(
        json.dumps(
            {
                "generated_at": utc_now_iso(),
                "table_preferences": {
                    "max_tables_per_node": 1,
                    "only_when_structured": True,
                },
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return target
