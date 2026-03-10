"""Word export agent that renders layout blocks into output.docx."""

from __future__ import annotations

import json
import tempfile
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor
from lxml import etree
from PIL import Image


class WordExportAgent:
    """Render layout blocks into a docx based on the standard template."""

    def export(
        self,
        *,
        template_path: Path,
        layout_blocks_path: Path,
        output_path: Path,
    ) -> tuple[Path, list[str]]:
        if not template_path.exists():
            raise FileNotFoundError(f"Template missing: {template_path}")

        try:
            document = Document(template_path)
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError(f"Failed to open template: {template_path}") from exc

        _remove_initial_empty_paragraph(document)
        payload = json.loads(layout_blocks_path.read_text(encoding="utf-8"))
        blocks = payload.get("blocks") or []
        warnings = list(payload.get("warnings") or [])
        writer = _DocumentWriter(document)

        for block in blocks:
            block_type = block.get("type")
            if block_type == "page_break":
                document.add_page_break()
                continue
            if block_type == "heading":
                writer.add_paragraph(
                    text=str(block.get("text") or ""),
                    style_name=str(block.get("style_name") or "Normal"),
                    warnings=warnings,
                )
                continue
            if block_type == "paragraph":
                writer.add_body_paragraph(
                    text=str(block.get("text") or ""),
                    style_name=str(block.get("style_name") or "Normal"),
                    warnings=warnings,
                )
                continue
            if block_type == "section_heading":
                writer.add_section_heading(
                    text=str(block.get("text") or ""),
                    style_name=str(block.get("style_name") or "Normal"),
                    warnings=warnings,
                )
                continue
            if block_type == "table":
                writer.add_table(block=block, warnings=warnings)
                continue
            if block_type == "image":
                writer.add_image(block=block, warnings=warnings)
                continue
            warnings.append(f"Unsupported layout block type: {block_type}")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        numbering_warnings = _normalize_docx_numbering(output_path)
        warnings.extend(numbering_warnings)
        return output_path, warnings


class _DocumentWriter:
    """Append content to a docx while preserving template styles."""

    DEFAULT_IMAGE_ASPECT_RATIO = "3:2"
    IMAGE_ASPECT_RATIO_VALUES = {
        "2:1": 2.0,
        "3:2": 1.5,
    }
    STYLE_ALIASES = {
        "标题 1": ("Heading 1",),
        "标题 2": ("Heading 2",),
        "标题 3": ("Heading 3",),
        "标题 4": ("Heading 4",),
        "Heading 1": ("标题 1",),
        "Heading 2": ("标题 2",),
        "Heading 3": ("标题 3",),
        "Heading 4": ("标题 4",),
    }

    def __init__(self, document: DocumentObject) -> None:
        self.document = document

    def add_paragraph(
        self,
        *,
        text: str,
        style_name: str,
        warnings: list[str],
    ) -> None:
        paragraph = self._next_paragraph()
        paragraph.text = text
        resolved_style = self._resolve_style(
            style_name=style_name,
            fallback="Normal",
            warnings=warnings,
        )
        paragraph.style = resolved_style
        if not self._is_heading_style(resolved_style):
            _disable_list_numbering(paragraph)
        paragraph.paragraph_format.keep_with_next = True

    def add_section_heading(
        self,
        *,
        text: str,
        style_name: str,
        warnings: list[str],
    ) -> None:
        paragraph = self._next_paragraph()
        paragraph.style = self._resolve_style(
            style_name=style_name,
            fallback="Normal",
            warnings=warnings,
        )
        _disable_list_numbering(paragraph)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    def add_body_paragraph(
        self,
        *,
        text: str,
        style_name: str,
        warnings: list[str],
    ) -> None:
        paragraph = self._next_paragraph()
        paragraph.text = text
        paragraph.style = self._resolve_style(
            style_name=style_name,
            fallback="Normal",
            warnings=warnings,
        )
        _disable_list_numbering(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    def add_table(self, *, block: dict[str, Any], warnings: list[str]) -> None:
        headers = [str(item) for item in (block.get("headers") or [])]
        rows = block.get("rows") or []
        if not headers:
            warnings.append(f"Table {block.get('table_id')} has no headers; skipped.")
            return

        title = str(block.get("title") or "").strip()
        if title:
            self.add_paragraph(text=title, style_name="Normal", warnings=warnings)

        table = self.document.add_table(rows=1, cols=len(headers))
        style_name = str(block.get("style_name") or "BiddingTable")
        if style_name in {style.name for style in self.document.styles}:
            table.style = style_name
        else:
            warnings.append(f"Table style {style_name} missing in template; using default table style.")

        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_cells[idx].text = header
            self._center_table_cell(header_cells[idx])

        for row in rows:
            cells = table.add_row().cells
            values = [str(item) for item in row]
            for idx in range(len(headers)):
                cells[idx].text = values[idx] if idx < len(values) else ""
                self._center_table_cell(cells[idx])

    def add_image(self, *, block: dict[str, Any], warnings: list[str]) -> None:
        image_path = Path(str(block.get("path") or ""))
        if not image_path.exists():
            warnings.append(f"Image file missing: {image_path}")
            return

        aspect_ratio = self._normalize_image_aspect_ratio(block.get("aspect_ratio"))
        render_path, cleanup_path = self._prepare_image_for_layout(
            image_path=image_path,
            aspect_ratio=aspect_ratio,
            warnings=warnings,
        )
        width, height = self._fit_image_size(aspect_ratio=aspect_ratio)
        paragraph = self._next_paragraph()
        paragraph.style = self._resolve_style(
            style_name="Normal",
            fallback="Normal",
            warnings=warnings,
        )
        _disable_list_numbering(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run()
        try:
            run.add_picture(str(render_path), width=width, height=height)
        finally:
            if cleanup_path is not None and cleanup_path.exists():
                cleanup_path.unlink()

    def _next_paragraph(self):
        return self.document.add_paragraph()

    def _resolve_style(
        self,
        *,
        style_name: str,
        fallback: str,
        warnings: list[str],
    ) -> str | None:
        style_names = {style.name for style in self.document.styles}
        if style_name in style_names:
            return style_name

        for alias in self.STYLE_ALIASES.get(style_name, ()):
            if alias in style_names:
                warnings.append(
                    f"Style {style_name} missing in template; fallback to {alias}."
                )
                return alias

        level = self._heading_level(style_name)
        if level is not None:
            while level >= 1:
                for candidate in (f"标题 {level}", f"Heading {level}"):
                    if candidate in style_names:
                        warnings.append(
                            f"Style {style_name} missing in template; fallback to {candidate}."
                        )
                        return candidate
                level -= 1

        if fallback in style_names:
            warnings.append(f"Style {style_name} missing in template; fallback to {fallback}.")
            return fallback
        warnings.append(f"Style {style_name} missing and no fallback found; document default used.")
        return None

    def _fit_image_size(self, *, aspect_ratio: str) -> tuple[Emu, Emu]:
        section = self.document.sections[-1]
        available_width = int(section.page_width - section.left_margin - section.right_margin)
        available_height = int(section.page_height - section.top_margin - section.bottom_margin)
        aspect_ratio_value = self.IMAGE_ASPECT_RATIO_VALUES[
            self._normalize_image_aspect_ratio(aspect_ratio)
        ]
        target_width = available_width
        target_height = int(target_width / aspect_ratio_value)
        if target_height > available_height:
            target_height = available_height
            target_width = int(target_height * aspect_ratio_value)
        return Emu(target_width), Emu(target_height)

    def _prepare_image_for_layout(
        self,
        *,
        image_path: Path,
        aspect_ratio: str,
        warnings: list[str],
    ) -> tuple[Path, Path | None]:
        try:
            with Image.open(image_path) as source:
                source_rgb = source.convert("RGB")
                target_ratio = self.IMAGE_ASPECT_RATIO_VALUES[
                    self._normalize_image_aspect_ratio(aspect_ratio)
                ]
                source_ratio = source_rgb.width / max(source_rgb.height, 1)
                if abs(source_ratio - target_ratio) < 0.01:
                    return image_path, None
                if source_ratio > target_ratio:
                    canvas_width = source_rgb.width
                    canvas_height = max(
                        source_rgb.height,
                        int(round(source_rgb.width / target_ratio)),
                    )
                else:
                    canvas_height = source_rgb.height
                    canvas_width = max(
                        source_rgb.width,
                        int(round(source_rgb.height * target_ratio)),
                    )
                canvas = Image.new("RGB", (canvas_width, canvas_height), (255, 255, 255))
                offset = (
                    (canvas_width - source_rgb.width) // 2,
                    (canvas_height - source_rgb.height) // 2,
                )
                canvas.paste(source_rgb, offset)
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"Failed to normalize image canvas for {image_path}: {exc}")
            return image_path, None

        handle = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        temp_path = Path(handle.name)
        handle.close()
        canvas.save(temp_path)
        return temp_path, temp_path

    @staticmethod
    def _center_table_cell(cell) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    @staticmethod
    def _heading_level(style_name: str | None) -> int | None:
        if not style_name:
            return None
        normalized = style_name.strip()
        for prefix in ("标题 ", "Heading "):
            if normalized.startswith(prefix):
                suffix = normalized[len(prefix) :].strip()
                if suffix.isdigit():
                    return int(suffix)
        return None

    @classmethod
    def _is_heading_style(cls, style_name: str | None) -> bool:
        return cls._heading_level(style_name) is not None

    @classmethod
    def _normalize_image_aspect_ratio(cls, aspect_ratio: Any) -> str:
        normalized = str(aspect_ratio or "").strip()
        if normalized in cls.IMAGE_ASPECT_RATIO_VALUES:
            return normalized
        return cls.DEFAULT_IMAGE_ASPECT_RATIO


def _remove_initial_empty_paragraph(document: DocumentObject) -> None:
    if len(document.paragraphs) != 1:
        return
    paragraph = document.paragraphs[0]
    if paragraph.text.strip():
        return
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _disable_list_numbering(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for child in list(ppr):
        if child.tag == qn("w:numPr"):
            ppr.remove(child)


def _normalize_docx_numbering(docx_path: Path) -> list[str]:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    heading_style_names = {
        f"heading {level}" for level in range(1, 5)
    } | {
        f"标题 {level}" for level in range(1, 5)
    }
    warnings: list[str] = []

    with ZipFile(docx_path, "r") as source_zip:
        styles_xml = source_zip.read("word/styles.xml")
        document_xml = source_zip.read("word/document.xml")
        other_files = {
            info.filename: source_zip.read(info.filename)
            for info in source_zip.infolist()
            if info.filename not in {"word/styles.xml", "word/document.xml"}
        }

    styles_root = etree.fromstring(styles_xml)
    document_root = etree.fromstring(document_xml)

    heading_style_ids: set[str] = set()
    normal_style_id: str | None = None
    for style_node in styles_root.xpath("//w:style[@w:type='paragraph']", namespaces=namespace):
        name_nodes = style_node.xpath("./w:name", namespaces=namespace)
        if not name_nodes:
            continue
        style_name = str(name_nodes[0].get(qn("w:val")) or "").strip().lower()
        if style_name not in heading_style_names:
            continue
        style_id = style_node.get(qn("w:styleId"))
        if not style_id:
            warnings.append(f"Heading style {style_name} missing styleId.")
            continue
        heading_style_ids.add(style_id)

    if not heading_style_ids:
        warnings.append("Heading styles 1-4 missing in template styles.xml.")

    normal_style_nodes = styles_root.xpath(
        "//w:style[w:name[@w:val='Normal']]",
        namespaces=namespace,
    )
    if normal_style_nodes:
        normal_style_id = normal_style_nodes[0].get(qn("w:styleId"))

    for paragraph in document_root.xpath("//w:p", namespaces=namespace):
        ppr_nodes = paragraph.xpath("./w:pPr", namespaces=namespace)
        if ppr_nodes:
            ppr = ppr_nodes[0]
        else:
            ppr = etree.Element(qn("w:pPr"))
            paragraph.insert(0, ppr)
        style_nodes = ppr.xpath("./w:pStyle", namespaces=namespace)
        style_id = style_nodes[0].get(qn("w:val")) if style_nodes else None
        for numpr in ppr.xpath("./w:numPr", namespaces=namespace):
            ppr.remove(numpr)
        if style_id in heading_style_ids:
            continue
        if (
            normal_style_id
            and style_id is None
            and not paragraph.xpath("ancestor::w:tc", namespaces=namespace)
        ):
            pstyle = etree.Element(qn("w:pStyle"))
            pstyle.set(qn("w:val"), normal_style_id)
            ppr.insert(0, pstyle)

    normalized_xml = etree.tostring(
        document_root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone="yes",
    )
    residual_numid_zero = normalized_xml.count(b'w:numId w:val="0"')
    if residual_numid_zero:
        warnings.append(f"Residual numId=0 found after normalization: {residual_numid_zero}")

    temp_output = docx_path.with_suffix(".tmp.docx")
    with ZipFile(temp_output, "w", compression=ZIP_DEFLATED) as target_zip:
        target_zip.writestr("word/styles.xml", styles_xml)
        target_zip.writestr("word/document.xml", normalized_xml)
        for filename, content in other_files.items():
            target_zip.writestr(filename, content)

    temp_output.replace(docx_path)
    return warnings
