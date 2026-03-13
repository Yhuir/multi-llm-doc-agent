"""Chunked requirement parser that extracts bidding requirements from the full doc."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib import error as urlerror
from urllib import request as urlrequest

from docx import Document

from backend.models.enums import AgentResult
from backend.models.schemas import (
    RequirementConstraints,
    RequirementDocument,
    RequirementItem,
    RequirementProject,
    RequirementScope,
    RequirementSubsystem,
    SourceIndexItem,
    utc_now_iso,
)


@dataclass
class _ClauseCandidate:
    source_ref: str
    line_no: int
    raw_text: str
    clause_code: str
    title: str
    normalized_key: str


@dataclass
class _SubsystemBucket:
    name: str
    description_lines: list[str] = field(default_factory=list)
    requirements: list[RequirementItem] = field(default_factory=list)
    interfaces: list[str] = field(default_factory=list)


@dataclass
class _ParagraphEntry:
    line_no: int
    text: str
    paragraph_id: str
    heading_level: int | None = None


@dataclass
class _ExtractedSubsystem:
    name: str
    description: str
    source_refs: list[str] = field(default_factory=list)


@dataclass
class _ChunkExtraction:
    overview_points: list[str] = field(default_factory=list)
    requirements: list[RequirementItem] = field(default_factory=list)
    subsystems: list[_ExtractedSubsystem] = field(default_factory=list)
    standards: list[RequirementItem] = field(default_factory=list)
    acceptance: list[RequirementItem] = field(default_factory=list)


@dataclass
class _RequirementConsolidation:
    project_name: str
    overview: str
    subsystems: list[_ExtractedSubsystem] = field(default_factory=list)
    standards: list[str] = field(default_factory=list)
    acceptance: list[str] = field(default_factory=list)


class RequirementParserAgent:
    """Parse docx into requirement.json by chunking the full document and using the text model."""

    _MINIMAX_TEXT_ENDPOINT = "https://api.minimaxi.com/v1/text/chatcompletion_v2"
    _WHATAI_TEXT_ENDPOINT = "https://api.whatai.cc/v1/chat/completions"
    _JSON_BLOCK_PATTERN = re.compile(r"```(?:json)?\s*(\{.*?\})\s*```", flags=re.S)
    _CLAUSE_PATTERN = re.compile(r"^(?P<code>\d+(?:\.\d+){1,5})[\s.、-]*(?P<body>.+)$")
    _NUMBERED_HEADING_PATTERN = re.compile(r"^(?P<code>\d+(?:\.\d+){0,3})[\s、.．]+(?P<body>.+)$")
    _CHAPTER_PATTERN = re.compile(r"^第[一二三四五六七八九十百千0-9]+[章节篇部分]\s*(?P<title>.*)$")
    _LIST_PATTERN = re.compile(r"^[（(]?[一二三四五六七八九十0-9]+[)）.、]\s*(?P<body>.+)$")
    _STANDARD_PATTERN = re.compile(
        r"(?<![A-Za-z0-9])(?:GB|JGJ|DL/T|IEC|ISO|T/[A-Za-z0-9]+)[A-Za-z0-9\-/]*(?![A-Za-z0-9])",
        re.IGNORECASE,
    )
    _SOURCE_REF_PATTERN = re.compile(r"^p\d+#L\d+$")
    _SUBSYSTEM_PATTERN = re.compile(
        r"([A-Za-z0-9\u4e00-\u9fff（）()\-]{2,32}?"
        r"(?:控制系统|回收系统|管理系统|供热系统|热泵系统|子系统|平台|模块|装置|系统))"
    )
    _SYSTEM_KEYWORDS = ("系统", "平台", "子系统", "模块", "控制", "装置", "网络", "中心")
    _GENERIC_HEADINGS = (
        "服务要求",
        "技术要求",
        "技术规范",
        "技术方案",
        "建设内容",
        "功能要求",
        "项目概述",
        "总体要求",
    )
    _NOISE_PREFIXES = (
        "根据生产管理的要求，",
        "根据要求，",
        "根据生产管理要求，",
        "针对",
        "要求对",
        "对",
    )
    _ACTION_PREFIXES = (
        "优化和完善",
        "优化完善",
        "优化升级",
        "升级改造",
        "优化",
        "完善",
        "改造",
        "升级",
        "开发",
        "建设",
        "实施",
        "新建",
        "扩建",
        "调整",
        "提升",
    )
    _CONTEXT_PREFIXES = ("当前的", "当前", "现有的", "现有", "原有的", "原有", "本次", "针对")
    _REQUIREMENT_KEYWORDS = (
        "应",
        "须",
        "必须",
        "不得",
        "负责",
        "要求",
        "提供",
        "满足",
        "完成",
        "实现",
        "采用",
        "具备",
        "接口",
        "验收",
        "调试",
        "施工",
        "培训",
        "售后",
        "维保",
        "质保",
        "交付",
        "服务",
        "安装",
        "控制",
        "监控",
        "网络",
        "通讯",
        "响应",
        "整改",
        "升级",
        "改造",
    )
    _REQUIREMENT_TYPES = {
        "scope",
        "technical",
        "service",
        "acceptance",
        "standard",
        "interface",
        "delivery",
        "schedule",
        "quality",
        "training",
        "warranty",
        "safety",
        "procurement",
        "operation",
        "maintenance",
        "other",
        "verbatim_requirement",
    }

    def parse(
        self,
        *,
        task_id: str,
        upload_file_path: Path,
        fallback_title: str,
        generation_config: dict[str, Any] | None = None,
    ) -> tuple[RequirementDocument, dict]:
        if upload_file_path.suffix.lower() != ".docx":
            raise ValueError("Only .docx is supported.")

        provider, model_name, api_key = self._resolve_generation_config(generation_config)
        paragraphs = self._extract_docx_paragraphs(upload_file_path)
        lines = [paragraph.text for paragraph in paragraphs]
        has_original_content = bool(lines)
        if not lines:
            lines = [fallback_title.strip() or "项目需求说明"]
            paragraphs = [
                _ParagraphEntry(
                    line_no=1,
                    text=lines[0],
                    paragraph_id="para_1",
                    heading_level=None,
                )
            ]

        source_index = self._build_source_index(paragraphs)
        chunks = self._build_chunks(paragraphs)
        chunk_extractions = [
            self._extract_chunk_payload(
                chunk=chunk,
                provider=provider,
                model_name=model_name,
                api_key=api_key,
                focus_refs=[],
            )
            for chunk in chunks
        ]
        aggregate = self._aggregate_chunk_extractions(chunk_extractions, source_index=source_index)

        uncovered_refs = self._find_uncovered_requirement_refs(
            paragraphs=paragraphs,
            covered_refs={item.source_ref for item in aggregate.requirements if item.source_ref},
        )
        if uncovered_refs:
            for focused_chunk in self._build_chunks_from_refs(paragraphs, uncovered_refs):
                follow_up = self._extract_chunk_payload(
                    chunk=focused_chunk,
                    provider=provider,
                    model_name=model_name,
                    api_key=api_key,
                    focus_refs=[f"p1#L{entry.line_no}" for entry in focused_chunk],
                )
                aggregate = self._aggregate_chunk_extractions(
                    [aggregate, follow_up],
                    source_index=source_index,
                )

        uncovered_refs = self._find_uncovered_requirement_refs(
            paragraphs=paragraphs,
            covered_refs={item.source_ref for item in aggregate.requirements if item.source_ref},
        )
        coverage_closure_items: list[RequirementItem] = []
        if uncovered_refs:
            coverage_closure_items = self._build_verbatim_closure_requirements(
                paragraphs=paragraphs,
                refs=uncovered_refs,
            )
            aggregate = self._aggregate_chunk_extractions(
                [
                    aggregate,
                    _ChunkExtraction(requirements=coverage_closure_items),
                ],
                source_index=source_index,
            )

        consolidation = self._consolidate_extractions(
            aggregate=aggregate,
            source_index=source_index,
            fallback_title=fallback_title,
            provider=provider,
            model_name=model_name,
            api_key=api_key,
        )
        subsystems = self._materialize_subsystems(
            subsystem_candidates=consolidation.subsystems or aggregate.subsystems,
            bidding_requirements=aggregate.requirements,
        )
        project_name = consolidation.project_name
        overview = consolidation.overview
        standards = consolidation.standards or [item.value for item in aggregate.standards]
        acceptance = consolidation.acceptance or [item.value for item in aggregate.acceptance]

        requirement = RequirementDocument(
            project=RequirementProject(
                name=project_name,
                customer="",
                location="",
                duration_days=None,
                milestones=[],
            ),
            scope=RequirementScope(
                overview=overview,
                subsystems=subsystems,
            ),
            constraints=RequirementConstraints(
                standards=standards,
                acceptance=acceptance,
            ),
            bidding_requirements=aggregate.requirements,
            source_index=source_index,
        )

        missing_fields: list[str] = []
        warnings: list[str] = []
        if not has_original_content:
            missing_fields.extend(["project.name", "scope.overview", "scope.subsystems"])
        if not standards:
            missing_fields.append("constraints.standards")
        if not acceptance:
            missing_fields.append("constraints.acceptance")
        if not subsystems:
            warnings.append("No subsystem could be consolidated from the full-document extraction result.")
        if coverage_closure_items:
            warnings.append(
                f"Coverage closure added {len(coverage_closure_items)} verbatim requirement items for uncovered refs."
            )

        parse_report = {
            "task_id": task_id,
            "source_file": str(upload_file_path),
            "result": AgentResult.PASS.value,
            "paragraph_count": len(lines),
            "subsystem_count": len(subsystems),
            "chunk_count": len(chunks),
            "bidding_requirement_count": len(requirement.bidding_requirements),
            "coverage_closure_count": len(coverage_closure_items),
            "missing_fields": missing_fields,
            "warnings": warnings,
            "generated_at": utc_now_iso(),
        }
        return requirement, parse_report

    def _resolve_generation_config(self, generation_config: dict[str, Any] | None) -> tuple[str, str, str]:
        config = generation_config or {}
        provider = str(config.get("text_provider") or "").strip().lower()
        model_name = str(config.get("text_model_name") or "MiniMax-M2.5").strip() or "MiniMax-M2.5"
        api_key = str(config.get("text_api_key") or "").strip()
        if provider not in {"minimax", "whatai", "google"}:
            raise ValueError(
                "Requirement parsing requires a supported text model provider. "
                f"Current provider: {provider or 'unset'}."
            )
        if not api_key:
            raise ValueError("Requirement parsing requires `text_api_key`. Please configure the text model first.")
        return provider, model_name, api_key

    def _request_minimax_completion(self, *, api_key: str, model_name: str, prompt: str) -> str:
        payload = {
            "model": model_name,
            "messages": [
                {
                    "role": "system",
                    "name": "Requirement Parser",
                    "content": (
                        "你是招标需求全文解析器。"
                        "必须通读输入文本，提取其中所有招标要求，禁止套用模板或遗漏要求。"
                    ),
                },
                {"role": "user", "name": "user", "content": prompt},
            ],
            "temperature": 0.1,
        }
        req = urlrequest.Request(
            self._MINIMAX_TEXT_ENDPOINT,
            data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            },
            method="POST",
        )

        try:
            with urlrequest.urlopen(req, timeout=240) as response:
                body = response.read().decode("utf-8")
        except urlerror.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="ignore")
            raise ValueError(f"Requirement parsing model request failed: HTTP {exc.code} {detail}") from exc
        except (urlerror.URLError, TimeoutError) as exc:
            raise ValueError(f"Requirement parsing model request failed: {exc}") from exc

        try:
            data = json.loads(body)
        except json.JSONDecodeError as exc:
            raise ValueError("Requirement parsing model returned invalid JSON payload.") from exc

        base_resp = data.get("base_resp") or {}
        if base_resp.get("status_code") not in (None, 0):
            raise ValueError(
                f"Requirement parsing model returned error: {base_resp.get('status_msg') or base_resp.get('status_code')}"
            )

        content = (
            ((data.get("choices") or [{}])[0].get("message") or {}).get("content")
            if isinstance(data.get("choices"), list)
            else None
        )
        if not isinstance(content, str) or not content.strip():
            raise ValueError("Requirement parsing model returned empty content.")
        return content

    def _request_whatai_completion(self, *, api_key: str, model_name: str, prompt: str) -> str:
        payload = {
            "model": model_name,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "你是招标需求全文解析器。"
                        "必须通读输入文本，提取其中所有招标要求，禁止套用模板或遗漏要求。"
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.1,
        }
        req = urlrequest.Request(
            self._WHATAI_TEXT_ENDPOINT,
            data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            },
            method="POST",
        )

        try:
            with urlrequest.urlopen(req, timeout=240) as response:
                body = response.read().decode("utf-8")
        except urlerror.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="ignore")
            raise ValueError(f"Requirement parsing model request failed: HTTP {exc.code} {detail}") from exc
        except (urlerror.URLError, TimeoutError) as exc:
            raise ValueError(f"Requirement parsing model request failed: {exc}") from exc

        try:
            data = json.loads(body)
        except json.JSONDecodeError as exc:
            raise ValueError("Requirement parsing model returned invalid JSON payload.") from exc

        if data.get("error"):
            raise ValueError(f"Requirement parsing model returned error: {data['error']}")

        content = (
            ((data.get("choices") or [{}])[0].get("message") or {}).get("content")
            if isinstance(data.get("choices"), list)
            else None
        )
        if not isinstance(content, str) or not content.strip():
            raise ValueError("Requirement parsing model returned empty content.")
        return content

    def _build_chunks(
        self,
        paragraphs: list[_ParagraphEntry],
        *,
        max_chars: int = 5200,
        max_paragraphs: int = 32,
    ) -> list[list[_ParagraphEntry]]:
        chunks: list[list[_ParagraphEntry]] = []
        current: list[_ParagraphEntry] = []
        current_chars = 0
        for entry in paragraphs:
            estimated_chars = len(entry.text) + 24
            if current and (len(current) >= max_paragraphs or current_chars + estimated_chars > max_chars):
                chunks.append(current)
                current = []
                current_chars = 0
            current.append(entry)
            current_chars += estimated_chars
        if current:
            chunks.append(current)
        return chunks

    def _build_chunks_from_refs(
        self,
        paragraphs: list[_ParagraphEntry],
        refs: list[str],
        *,
        max_chars: int = 3600,
        max_paragraphs: int = 18,
    ) -> list[list[_ParagraphEntry]]:
        ref_set = set(refs)
        focused = [entry for entry in paragraphs if f"p1#L{entry.line_no}" in ref_set]
        return self._build_chunks(focused, max_chars=max_chars, max_paragraphs=max_paragraphs)

    def _extract_chunk_payload(
        self,
        *,
        chunk: list[_ParagraphEntry],
        provider: str,
        model_name: str,
        api_key: str,
        focus_refs: list[str],
    ) -> _ChunkExtraction:
        prompt = self._build_chunk_prompt(chunk=chunk, focus_refs=focus_refs)
        request_fn = self._request_minimax_completion if provider == "minimax" else self._request_whatai_completion
        content = request_fn(api_key=api_key, model_name=model_name, prompt=prompt)
        last_error: ValueError | None = None
        for attempt in range(2):
            try:
                return self._parse_chunk_output(content=content, chunk=chunk)
            except ValueError as exc:
                last_error = exc
                if attempt >= 1:
                    break
                repair_prompt = "\n".join(
                    [
                        prompt,
                        "",
                        "上一次输出存在问题：",
                        str(exc),
                        "",
                        "请修正并只输出合法 JSON。上一次无效输出如下：",
                        content,
                    ]
                )
                content = request_fn(api_key=api_key, model_name=model_name, prompt=repair_prompt)
        assert last_error is not None
        raise last_error

    def _build_chunk_prompt(self, *, chunk: list[_ParagraphEntry], focus_refs: list[str]) -> str:
        source_lines = [f"p1#L{entry.line_no} | {entry.text}" for entry in chunk]
        focus_lines = [f"- {ref}" for ref in focus_refs]
        return "\n".join(
            [
                "任务：阅读以下招标文档片段，提取其中所有招标要求。",
                "要求：",
                "1. 必须逐条检查片段中的段落和小标题，只要包含招标范围、技术要求、实施要求、接口要求、供货要求、服务要求、培训要求、质保要求、验收要求、进度要求或标准约束，就必须提取。",
                "2. 禁止遗漏片段中的要求，也禁止编造文档中没有出现的要求。",
                "3. 每条 requirements 必须给出 source_ref，且 source_ref 必须来自下方片段中的引用键。",
                "4. subsystems 只保留文档中明确出现的系统/子系统/平台/模块名称。",
                "5. standards 只保留文档中明确出现的标准号。",
                "6. acceptance 只保留文档中明确出现的验收、试运行、交付、签认、资料提交相关要求。",
                "7. 只输出 JSON，不要输出解释。",
                "",
                "输出 JSON 格式：",
                "{",
                '  "overview_points": ["片段中的关键范围或目标摘要"],',
                '  "requirements": [',
                '    {"type": "technical", "key": "短键", "value": "提炼后的要求内容", "source_ref": "p1#L10"}',
                "  ],",
                '  "subsystems": [',
                '    {"name": "系统名称", "description": "该系统在片段中的要求摘要", "source_refs": ["p1#L10"]}',
                "  ],",
                '  "standards": [',
                '    {"name": "GB50348", "source_ref": "p1#L11"}',
                "  ],",
                '  "acceptance": [',
                '    {"value": "完成试运行并提交验收资料", "source_ref": "p1#L12"}',
                "  ]",
                "}",
                "",
                "必须优先复查的上轮未覆盖引用：",
                *(focus_lines or ["- 无"]),
                "",
                "文档片段：",
                *source_lines,
            ]
        )

    def _parse_chunk_output(self, *, content: str, chunk: list[_ParagraphEntry]) -> _ChunkExtraction:
        payload = self._extract_json_payload(content)
        allowed_refs = {f"p1#L{entry.line_no}" for entry in chunk}
        paragraph_lookup = {f"p1#L{entry.line_no}": entry.text for entry in chunk}

        overview_points = self._normalize_string_list(payload.get("overview_points"))
        requirements = self._parse_requirement_items(
            payload.get("requirements"),
            allowed_refs=allowed_refs,
            paragraph_lookup=paragraph_lookup,
        )
        standards = self._parse_named_items(
            payload.get("standards"),
            item_type="standard",
            allowed_refs=allowed_refs,
            paragraph_lookup=paragraph_lookup,
        )
        acceptance = self._parse_named_items(
            payload.get("acceptance"),
            item_type="acceptance",
            allowed_refs=allowed_refs,
            paragraph_lookup=paragraph_lookup,
        )
        subsystems = self._parse_subsystem_items(
            payload.get("subsystems"),
            allowed_refs=allowed_refs,
            paragraph_lookup=paragraph_lookup,
        )
        return _ChunkExtraction(
            overview_points=overview_points,
            requirements=requirements,
            subsystems=subsystems,
            standards=standards,
            acceptance=acceptance,
        )

    def _parse_requirement_items(
        self,
        raw_items: Any,
        *,
        allowed_refs: set[str],
        paragraph_lookup: dict[str, str],
    ) -> list[RequirementItem]:
        if not isinstance(raw_items, list):
            return []
        items: list[RequirementItem] = []
        for index, raw in enumerate(raw_items, start=1):
            if not isinstance(raw, dict):
                continue
            value = self._normalize_text(str(raw.get("value") or raw.get("text") or ""))
            if not value:
                continue
            source_ref = self._resolve_source_ref(
                raw.get("source_ref"),
                text=value,
                allowed_refs=allowed_refs,
                paragraph_lookup=paragraph_lookup,
            )
            if source_ref is None:
                continue
            item_type = self._normalize_requirement_type(str(raw.get("type") or "other"))
            key = self._normalize_requirement_key(str(raw.get("key") or ""), fallback=value, index=index)
            items.append(
                RequirementItem(
                    type=item_type,
                    key=key,
                    value=value,
                    source_ref=source_ref,
                )
            )
        return self._dedupe_requirement_items(items)

    def _parse_named_items(
        self,
        raw_items: Any,
        *,
        item_type: str,
        allowed_refs: set[str],
        paragraph_lookup: dict[str, str],
    ) -> list[RequirementItem]:
        if not isinstance(raw_items, list):
            return []
        parsed: list[RequirementItem] = []
        for index, raw in enumerate(raw_items, start=1):
            if isinstance(raw, dict):
                value = self._normalize_text(str(raw.get("name") or raw.get("value") or raw.get("text") or ""))
                raw_ref = raw.get("source_ref")
            else:
                value = self._normalize_text(str(raw))
                raw_ref = None
            if not value:
                continue
            source_ref = self._resolve_source_ref(
                raw_ref,
                text=value,
                allowed_refs=allowed_refs,
                paragraph_lookup=paragraph_lookup,
            )
            if source_ref is None:
                continue
            parsed.append(
                RequirementItem(
                    type=item_type,
                    key=self._normalize_requirement_key("", fallback=value, index=index),
                    value=value,
                    source_ref=source_ref,
                )
            )
        return self._dedupe_requirement_items(parsed)

    def _parse_subsystem_items(
        self,
        raw_items: Any,
        *,
        allowed_refs: set[str],
        paragraph_lookup: dict[str, str],
    ) -> list[_ExtractedSubsystem]:
        if not isinstance(raw_items, list):
            return []
        results: list[_ExtractedSubsystem] = []
        for raw in raw_items:
            if not isinstance(raw, dict):
                continue
            name = self._normalize_subsystem_title(str(raw.get("name") or ""))
            description = self._normalize_text(str(raw.get("description") or ""))
            if not name:
                continue
            refs: list[str] = []
            raw_refs = raw.get("source_refs")
            if isinstance(raw_refs, list):
                for item in raw_refs:
                    source_ref = self._resolve_source_ref(
                        item,
                        text=description or name,
                        allowed_refs=allowed_refs,
                        paragraph_lookup=paragraph_lookup,
                    )
                    if source_ref and source_ref not in refs:
                        refs.append(source_ref)
            if not refs:
                inferred_ref = self._resolve_source_ref(
                    None,
                    text=description or name,
                    allowed_refs=allowed_refs,
                    paragraph_lookup=paragraph_lookup,
                )
                if inferred_ref:
                    refs.append(inferred_ref)
            results.append(_ExtractedSubsystem(name=name, description=description or name, source_refs=refs))
        deduped: list[_ExtractedSubsystem] = []
        seen: set[str] = set()
        for item in results:
            key = self._normalize_topic_key(item.name)
            if key in seen:
                continue
            seen.add(key)
            deduped.append(item)
        return deduped

    def _aggregate_chunk_extractions(
        self,
        extractions: list[_ChunkExtraction],
        *,
        source_index: dict[str, SourceIndexItem],
    ) -> _ChunkExtraction:
        overview_points: list[str] = []
        requirements: list[RequirementItem] = []
        standards: list[RequirementItem] = []
        acceptance: list[RequirementItem] = []
        subsystem_map: dict[str, _ExtractedSubsystem] = {}

        for extraction in extractions:
            for point in extraction.overview_points:
                if point and point not in overview_points:
                    overview_points.append(point)
            requirements.extend(extraction.requirements)
            standards.extend(extraction.standards)
            acceptance.extend(extraction.acceptance)
            for subsystem in extraction.subsystems:
                key = self._normalize_topic_key(subsystem.name)
                existing = subsystem_map.get(key)
                if existing is None:
                    subsystem_map[key] = _ExtractedSubsystem(
                        name=subsystem.name,
                        description=subsystem.description,
                        source_refs=list(subsystem.source_refs),
                    )
                    continue
                if subsystem.description and subsystem.description not in existing.description:
                    existing.description = "；".join(
                        [item for item in [existing.description, subsystem.description] if item][:2]
                    )
                for ref in subsystem.source_refs:
                    if ref not in existing.source_refs:
                        existing.source_refs.append(ref)

        requirements = self._dedupe_requirement_items(requirements)
        standards = self._dedupe_requirement_items(standards)
        acceptance = self._dedupe_requirement_items(acceptance)

        for item in [*standards, *acceptance]:
            if not any(
                existing.source_ref == item.source_ref and existing.value == item.value
                for existing in requirements
            ):
                requirements.append(item)
        requirements = self._sort_requirement_items(self._dedupe_requirement_items(requirements), source_index)

        return _ChunkExtraction(
            overview_points=overview_points,
            requirements=requirements,
            subsystems=list(subsystem_map.values()),
            standards=standards,
            acceptance=acceptance,
        )

    def _find_uncovered_requirement_refs(
        self,
        *,
        paragraphs: list[_ParagraphEntry],
        covered_refs: set[str],
    ) -> list[str]:
        uncovered: list[str] = []
        for entry in paragraphs:
            ref = f"p1#L{entry.line_no}"
            if ref in covered_refs:
                continue
            if self._is_requirement_candidate(entry):
                uncovered.append(ref)
        return uncovered

    def _is_requirement_candidate(self, entry: _ParagraphEntry) -> bool:
        text = self._normalize_text(entry.text)
        if not text:
            return False
        if entry.heading_level == 1 or self._CHAPTER_PATTERN.match(text):
            return False
        if entry.heading_level is not None and entry.heading_level >= 2:
            return True
        if self._CLAUSE_PATTERN.match(text):
            return True
        if self._STANDARD_PATTERN.search(text):
            return True
        if len(text) > 260:
            return False
        return any(keyword in text for keyword in self._REQUIREMENT_KEYWORDS)

    def _build_verbatim_closure_requirements(
        self,
        *,
        paragraphs: list[_ParagraphEntry],
        refs: list[str],
    ) -> list[RequirementItem]:
        lookup = {f"p1#L{entry.line_no}": entry for entry in paragraphs}
        items: list[RequirementItem] = []
        for index, ref in enumerate(refs, start=1):
            entry = lookup.get(ref)
            if entry is None:
                continue
            items.append(
                RequirementItem(
                    type="verbatim_requirement",
                    key=f"coverage_{index:03d}",
                    value=entry.text,
                    source_ref=ref,
                )
            )
        return items

    def _consolidate_extractions(
        self,
        *,
        aggregate: _ChunkExtraction,
        source_index: dict[str, SourceIndexItem],
        fallback_title: str,
        provider: str,
        model_name: str,
        api_key: str,
    ) -> _RequirementConsolidation:
        prompt = self._build_consolidation_prompt(
            aggregate=aggregate,
            source_index=source_index,
            fallback_title=fallback_title,
        )
        request_fn = self._request_minimax_completion if provider == "minimax" else self._request_whatai_completion
        content = request_fn(api_key=api_key, model_name=model_name, prompt=prompt)
        last_error: ValueError | None = None
        for attempt in range(2):
            try:
                return self._parse_consolidation_output(content=content, source_index=source_index)
            except ValueError as exc:
                last_error = exc
                if attempt >= 1:
                    break
                repair_prompt = "\n".join(
                    [
                        prompt,
                        "",
                        "上一次输出存在问题：",
                        str(exc),
                        "",
                        "请修正并只输出合法 JSON。上一次无效输出如下：",
                        content,
                    ]
                )
                content = request_fn(api_key=api_key, model_name=model_name, prompt=repair_prompt)
        assert last_error is not None
        raise last_error

    def _build_consolidation_prompt(
        self,
        *,
        aggregate: _ChunkExtraction,
        source_index: dict[str, SourceIndexItem],
        fallback_title: str,
    ) -> str:
        requirement_lines = [
            f"- {item.source_ref}: [{item.type}] {item.value}"
            for item in aggregate.requirements
            if item.source_ref and item.value
        ]
        subsystem_lines = [
            f"- {item.name}: {item.description} | refs={', '.join(item.source_refs) or '无'}"
            for item in aggregate.subsystems
            if item.name
        ]
        standard_lines = [f"- {item.value}" for item in aggregate.standards if item.value]
        acceptance_lines = [f"- {item.value}" for item in aggregate.acceptance if item.value]
        source_lines = [
            f"{ref} | {item.text}"
            for ref, item in sorted(
                source_index.items(),
                key=lambda pair: self._line_no_from_ref(pair[0]),
            )[:120]
        ]
        return "\n".join(
            [
                "任务：根据已分段提取出的全文招标要求，整编 requirement.json 的核心摘要字段。",
                "要求：",
                "1. project_name、overview、subsystems、standards、acceptance 都必须来自下方已提取的招标要求和原文索引概览。",
                "2. 禁止补充文档中不存在的系统、标准、验收结论或服务承诺。",
                "3. subsystems 只保留文档中明确出现并且对后续目录/正文有组织价值的系统或专题。",
                "4. 只输出 JSON。",
                "",
                "输出 JSON 格式：",
                "{",
                '  "project_name": "项目名称",',
                '  "overview": "全文范围概述",',
                '  "subsystems": [',
                '    {"name": "系统名称", "description": "系统要求摘要", "source_refs": ["p1#L10"]}',
                "  ],",
                '  "standards": ["GB50348"],',
                '  "acceptance": ["完成试运行并提交资料"]',
                "}",
                "",
                f"fallback_title：{self._normalize_text(fallback_title) or '未提供'}",
                "",
                "已提取的招标要求（必须优先使用）：",
                *(requirement_lines or ["- 无"]),
                "",
                "已识别系统候选：",
                *(subsystem_lines or ["- 无"]),
                "",
                "已识别标准：",
                *(standard_lines or ["- 无"]),
                "",
                "已识别验收/交付要求：",
                *(acceptance_lines or ["- 无"]),
                "",
                "原文索引概览（前 120 行，仅用于校核命名和范围）：",
                *source_lines,
            ]
        )

    def _parse_consolidation_output(
        self,
        *,
        content: str,
        source_index: dict[str, SourceIndexItem],
    ) -> _RequirementConsolidation:
        payload = self._extract_json_payload(content)
        project_name = self._normalize_text(str(payload.get("project_name") or payload.get("title") or ""))
        overview = self._normalize_text(str(payload.get("overview") or ""))
        if not project_name:
            raise ValueError("Requirement consolidation did not return project_name.")
        if not overview:
            raise ValueError("Requirement consolidation did not return overview.")
        allowed_refs = set(source_index.keys())
        paragraph_lookup = {ref: item.text for ref, item in source_index.items()}
        subsystems = self._parse_subsystem_items(
            payload.get("subsystems"),
            allowed_refs=allowed_refs,
            paragraph_lookup=paragraph_lookup,
        )
        standards = [item.value for item in self._parse_named_items(
            payload.get("standards"),
            item_type="standard",
            allowed_refs=allowed_refs,
            paragraph_lookup=paragraph_lookup,
        )]
        acceptance = [item.value for item in self._parse_named_items(
            payload.get("acceptance"),
            item_type="acceptance",
            allowed_refs=allowed_refs,
            paragraph_lookup=paragraph_lookup,
        )]
        return _RequirementConsolidation(
            project_name=project_name,
            overview=overview,
            subsystems=subsystems,
            standards=self._dedupe_strings(standards),
            acceptance=self._dedupe_strings(acceptance),
        )

    def _build_requirement_lookup(
        self,
        bidding_requirements: list[RequirementItem],
    ) -> dict[str, list[RequirementItem]]:
        lookup: dict[str, list[RequirementItem]] = {}
        for item in bidding_requirements:
            if not item.source_ref:
                continue
            lookup.setdefault(item.source_ref, []).append(item)
        return lookup

    def _materialize_subsystems(
        self,
        *,
        subsystem_candidates: list[_ExtractedSubsystem],
        bidding_requirements: list[RequirementItem],
    ) -> list[RequirementSubsystem]:
        requirement_lookup = self._build_requirement_lookup(bidding_requirements)
        subsystems: list[RequirementSubsystem] = []
        for candidate in subsystem_candidates:
            refs = [ref for ref in candidate.source_refs if ref]
            linked_requirements: list[RequirementItem] = []
            for ref in refs:
                for item in requirement_lookup.get(ref, []):
                    if item not in linked_requirements:
                        linked_requirements.append(item)
            if not linked_requirements:
                normalized_name = self._normalize_topic_key(candidate.name)
                for item in bidding_requirements:
                    if normalized_name and normalized_name in self._normalize_topic_key(item.value):
                        linked_requirements.append(item)
            subsystems.append(
                RequirementSubsystem(
                    name=candidate.name,
                    description=candidate.description or candidate.name,
                    requirements=self._sort_requirement_items(
                        self._dedupe_requirement_items(linked_requirements),
                        {item.source_ref: SourceIndexItem(page=1, paragraph_id="", text="") for item in linked_requirements if item.source_ref},
                    ),
                    interfaces=[item.value for item in linked_requirements if "接口" in item.value][:3],
                )
            )
        return subsystems

    def _extract_json_payload(self, content: str) -> dict[str, Any]:
        candidates = self._JSON_BLOCK_PATTERN.findall(content)
        first = content.find("{")
        last = content.rfind("}")
        if first != -1 and last != -1 and first < last:
            candidates.append(content[first : last + 1])
        for candidate in candidates:
            try:
                payload = json.loads(candidate)
            except json.JSONDecodeError:
                continue
            if isinstance(payload, dict):
                return payload
        raise ValueError("Requirement parsing model did not return parsable JSON.")

    def _normalize_string_list(self, raw_items: Any) -> list[str]:
        if not isinstance(raw_items, list):
            return []
        items: list[str] = []
        for raw in raw_items:
            value = self._normalize_text(str(raw))
            if value and value not in items:
                items.append(value)
        return items

    def _normalize_requirement_type(self, raw_type: str) -> str:
        normalized = self._normalize_text(raw_type).lower().replace(" ", "_")
        return normalized if normalized in self._REQUIREMENT_TYPES else "other"

    def _normalize_requirement_key(self, raw_key: str, *, fallback: str, index: int) -> str:
        normalized = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "_", self._normalize_text(raw_key)).strip("_")
        if not normalized:
            normalized = self._normalize_topic_key(fallback)
        return (normalized[:48] or f"item_{index:03d}").lower()

    def _resolve_source_ref(
        self,
        raw_ref: Any,
        *,
        text: str,
        allowed_refs: set[str],
        paragraph_lookup: dict[str, str],
    ) -> str | None:
        normalized_ref = self._normalize_text(str(raw_ref or ""))
        if normalized_ref in allowed_refs and self._SOURCE_REF_PATTERN.match(normalized_ref):
            return normalized_ref
        normalized_text = self._normalize_text(text)
        if normalized_text:
            exact_matches = [ref for ref, paragraph in paragraph_lookup.items() if normalized_text in paragraph]
            if len(exact_matches) == 1:
                return exact_matches[0]
        return None

    def _dedupe_requirement_items(self, items: list[RequirementItem]) -> list[RequirementItem]:
        deduped: list[RequirementItem] = []
        seen: set[tuple[str, str, str]] = set()
        for item in items:
            marker = (
                item.type,
                item.source_ref or "",
                self._normalize_text(item.value),
            )
            if marker in seen:
                continue
            seen.add(marker)
            deduped.append(item)
        return deduped

    def _sort_requirement_items(
        self,
        items: list[RequirementItem],
        source_index: dict[str, SourceIndexItem],
    ) -> list[RequirementItem]:
        return sorted(
            items,
            key=lambda item: (
                self._line_no_from_ref(item.source_ref or ""),
                self._normalize_text(item.value),
            ),
        )

    def _line_no_from_ref(self, source_ref: str) -> int:
        match = re.search(r"#L(\d+)$", source_ref)
        if match is None:
            return 10**9
        return int(match.group(1))

    def _dedupe_strings(self, values: list[str]) -> list[str]:
        deduped: list[str] = []
        for value in values:
            cleaned = self._normalize_text(value)
            if cleaned and cleaned not in deduped:
                deduped.append(cleaned)
        return deduped

    def _extract_docx_paragraphs(self, file_path: Path) -> list[_ParagraphEntry]:
        try:
            document = Document(file_path)
        except Exception as exc:
            raise ValueError(f"Cannot parse docx file: {file_path}") from exc

        entries: list[_ParagraphEntry] = []
        for paragraph in document.paragraphs:
            text = self._normalize_text(paragraph.text)
            if not text:
                continue
            style_name = self._normalize_text(getattr(getattr(paragraph, "style", None), "name", ""))
            heading_level = self._detect_heading_level(text=text, style_name=style_name)
            line_no = len(entries) + 1
            paragraph_id = (
                f"heading_l{heading_level}_{line_no}"
                if heading_level is not None
                else f"para_{line_no}"
            )
            entries.append(
                _ParagraphEntry(
                    line_no=line_no,
                    text=text,
                    paragraph_id=paragraph_id,
                    heading_level=heading_level,
                )
            )
        return entries

    def _build_source_index(self, paragraphs: list[_ParagraphEntry]) -> dict[str, SourceIndexItem]:
        source_index: dict[str, SourceIndexItem] = {}
        for paragraph in paragraphs:
            ref = f"p1#L{paragraph.line_no}"
            source_index[ref] = SourceIndexItem(
                page=1,
                paragraph_id=paragraph.paragraph_id,
                text=paragraph.text,
            )
        return source_index

    def _extract_subsystems(self, paragraphs: list[_ParagraphEntry]) -> list[RequirementSubsystem]:
        outline_subsystems = self._extract_outline_subsystems(paragraphs)
        if outline_subsystems:
            return outline_subsystems

        buckets: dict[str, _SubsystemBucket] = {}

        for paragraph in paragraphs:
            line = paragraph.text
            if not line:
                continue
            if paragraph.heading_level is not None:
                continue
            if self._is_generic_heading(line):
                continue

            candidate = self._build_clause_candidate(line=line, line_no=paragraph.line_no)
            if candidate is None:
                continue

            for expanded_title in self._expand_candidate_titles(candidate.raw_text, candidate.title):
                normalized_key = self._normalize_topic_key(expanded_title)
                bucket = buckets.setdefault(
                    normalized_key,
                    _SubsystemBucket(name=expanded_title),
                )
                if candidate.raw_text not in bucket.description_lines:
                    bucket.description_lines.append(candidate.raw_text)
                bucket.requirements.append(
                    RequirementItem(
                        type="clause" if candidate.clause_code else "heading",
                        key=candidate.clause_code or f"line_{candidate.line_no}",
                        value=candidate.raw_text,
                        source_ref=candidate.source_ref,
                    )
                )
                if "接口" in candidate.raw_text and candidate.raw_text not in bucket.interfaces:
                    bucket.interfaces.append(candidate.raw_text)

        if not buckets:
            return self._fallback_subsystems([paragraph.text for paragraph in paragraphs])

        subsystems: list[RequirementSubsystem] = []
        for bucket in buckets.values():
            subsystems.append(
                RequirementSubsystem(
                    name=bucket.name,
                    description="；".join(bucket.description_lines[:3]),
                    requirements=bucket.requirements,
                    interfaces=bucket.interfaces,
                )
            )
        return subsystems

    def _build_clause_candidate(self, *, line: str, line_no: int) -> _ClauseCandidate | None:
        clause_code = ""
        body = line

        while True:
            matched_clause = self._CLAUSE_PATTERN.match(body)
            if matched_clause is None:
                break
            clause_code = matched_clause.group("code")
            body = self._normalize_text(matched_clause.group("body"))

        chapter_match = self._CHAPTER_PATTERN.match(body)
        if chapter_match is not None:
            chapter_title = self._normalize_text(chapter_match.group("title"))
            body = chapter_title or body

        list_match = self._LIST_PATTERN.match(body)
        if list_match is not None:
            body = self._normalize_text(list_match.group("body"))

        title = self._derive_topic_title(body)
        if not self._looks_like_requirement_topic(title, body, clause_code):
            return None

        source_ref = f"p1#L{line_no}"
        return _ClauseCandidate(
            source_ref=source_ref,
            line_no=line_no,
            raw_text=line,
            clause_code=clause_code,
            title=title,
            normalized_key=self._normalize_topic_key(title),
        )

    def _fallback_subsystems(self, lines: list[str]) -> list[RequirementSubsystem]:
        fallback_lines = [
            self._normalize_text(item)
            for item in lines[1:8]
            if self._normalize_text(item) and not self._is_generic_heading(item)
        ]
        if not fallback_lines:
            fallback_lines = ["基础实施范围"]

        subsystems: list[RequirementSubsystem] = []
        for idx, line in enumerate(fallback_lines[:4], start=1):
            ref = f"p1#L{min(idx + 1, max(len(lines), 1))}"
            title = self._derive_topic_title(line)
            subsystems.append(
                RequirementSubsystem(
                    name=title or f"实施专题{idx}",
                    description=line,
                    requirements=[
                        RequirementItem(
                            type="fallback",
                            key=f"fallback_{idx}",
                            value=line,
                            source_ref=ref,
                        )
                    ],
                    interfaces=[line] if "接口" in line else [],
                )
            )
        return subsystems

    def _extract_outline_subsystems(self, paragraphs: list[_ParagraphEntry]) -> list[RequirementSubsystem]:
        headings = [entry for entry in paragraphs if entry.heading_level is not None and entry.heading_level <= 4]
        if not headings:
            return []

        if any(entry.heading_level == 2 for entry in headings):
            primary_level = 2
        elif any(entry.heading_level == 1 for entry in headings):
            primary_level = 1
        else:
            primary_level = min(entry.heading_level for entry in headings)

        primary_headings = [entry for entry in headings if entry.heading_level == primary_level]
        if primary_level == 1 and len(primary_headings) == 1 and any(entry.heading_level > 1 for entry in headings):
            primary_level = 2
            primary_headings = [entry for entry in headings if entry.heading_level == 2]

        if not primary_headings:
            return []

        subsystems: list[RequirementSubsystem] = []
        for index, heading in enumerate(primary_headings):
            next_line_no = (
                primary_headings[index + 1].line_no
                if index + 1 < len(primary_headings)
                else paragraphs[-1].line_no + 1
            )
            segment = [
                entry
                for entry in paragraphs
                if heading.line_no <= entry.line_no < next_line_no
            ]
            title = self._clean_heading_title(heading.text)
            if not title or self._is_generic_heading(title):
                continue

            requirements: list[RequirementItem] = [
                RequirementItem(
                    type=f"heading_l{heading.heading_level}",
                    key=heading.paragraph_id,
                    value=heading.text,
                    source_ref=f"p1#L{heading.line_no}",
                )
            ]
            description_parts: list[str] = []
            interfaces: list[str] = []

            for entry in segment[1:]:
                source_ref = f"p1#L{entry.line_no}"
                if entry.heading_level is not None and entry.heading_level > heading.heading_level:
                    requirements.append(
                        RequirementItem(
                            type=f"heading_l{entry.heading_level}",
                            key=entry.paragraph_id,
                            value=entry.text,
                            source_ref=source_ref,
                        )
                    )
                    description_parts.append(self._clean_heading_title(entry.text))
                    continue

                if entry.heading_level is None and len(entry.text) <= 120:
                    requirements.append(
                        RequirementItem(
                            type="text",
                            key=f"para_{entry.line_no}",
                            value=entry.text,
                            source_ref=source_ref,
                        )
                    )
                    if len(description_parts) < 3:
                        description_parts.append(entry.text)
                    if "接口" in entry.text and entry.text not in interfaces:
                        interfaces.append(entry.text)

                if len(requirements) >= 8:
                    break

            subsystems.append(
                RequirementSubsystem(
                    name=title,
                    description="；".join(description_parts[:3]) or title,
                    requirements=requirements,
                    interfaces=interfaces,
                )
            )

        return subsystems

    def _extract_project_name(
        self,
        paragraphs: list[_ParagraphEntry],
        subsystems: list[RequirementSubsystem],
        fallback_title: str,
    ) -> str:
        heading_level1 = [entry for entry in paragraphs if entry.heading_level == 1]
        if len(heading_level1) == 1:
            title = self._clean_heading_title(heading_level1[0].text)
            if title and not self._is_generic_heading(title):
                return title

        for entry in paragraphs[:20]:
            cleaned = self._normalize_text(entry.text)
            if not cleaned or self._is_generic_heading(cleaned):
                continue
            body = self._strip_leading_codes(cleaned)
            title = self._derive_topic_title(body)
            if 6 <= len(title) <= 28 and self._contains_topic_keyword(title):
                return title

        if subsystems:
            return subsystems[0].name

        fallback = self._normalize_text(fallback_title)
        if fallback:
            return fallback
        return "工程实施项目"

    def _extract_overview(self, paragraphs: list[_ParagraphEntry]) -> str:
        fragments: list[str] = []
        for entry in paragraphs[:16]:
            if entry.heading_level is not None:
                continue
            cleaned = self._normalize_text(entry.text)
            if not cleaned or self._is_generic_heading(cleaned):
                continue
            body = self._strip_leading_codes(cleaned)
            fragments.append(body)
            if len("；".join(fragments)) >= 80:
                break
        if fragments:
            return "；".join(fragments[:3])
        return "基于需求文件提取实施范围、约束条件和关键建设内容。"

    def _derive_topic_title(self, text: str) -> str:
        body = self._prepare_topic_body(text)
        body = re.split(r"[，。；：]", body, maxsplit=1)[0].strip(" -、:：;；")
        if "包括" in body and len(body) > 22:
            body = body.split("包括", 1)[0].strip()
        if "要求" in body and len(body) > 18 and not body.endswith("要求"):
            body = body.split("要求", 1)[0].strip()

        match = self._SUBSYSTEM_PATTERN.search(body)
        if match is not None:
            body = self._normalize_subsystem_title(match.group(1))

        body = body.strip(" -、:：;；")
        if len(body) > 32:
            body = body[:32].rstrip(" -、")
        return body or "实施专题"

    def _detect_heading_level(self, *, text: str, style_name: str) -> int | None:
        level_from_style = self._detect_heading_level_from_style(style_name)
        if level_from_style is not None:
            return level_from_style

        if self._CHAPTER_PATTERN.match(text):
            return 1

        numbered_match = self._NUMBERED_HEADING_PATTERN.match(text)
        if numbered_match is not None:
            code = numbered_match.group("code")
            body = self._normalize_text(numbered_match.group("body"))
            if self._looks_like_numbered_heading_body(body):
                return min(code.count(".") + 1, 4)

        return None

    def _detect_heading_level_from_style(self, style_name: str) -> int | None:
        normalized = style_name.lower().replace(" ", "")
        match = re.search(r"heading([1-4])", normalized)
        if match is not None:
            return int(match.group(1))
        match = re.search(r"标题([1-4])", style_name)
        if match is not None:
            return int(match.group(1))
        return None

    def _looks_like_numbered_heading_body(self, body: str) -> bool:
        cleaned = self._clean_heading_title(body)
        if not cleaned or len(cleaned) > 38:
            return False
        if body.endswith(("。", "；", ";")):
            return False
        punctuation_count = sum(body.count(token) for token in ("，", "。", "；", "：", ":", ",", ";"))
        if punctuation_count > 1:
            return False
        return True

    def _clean_heading_title(self, text: str) -> str:
        cleaned = self._normalize_text(text)
        chapter_match = self._CHAPTER_PATTERN.match(cleaned)
        if chapter_match is not None:
            cleaned = self._normalize_text(chapter_match.group("title")) or cleaned
        numbered_match = self._NUMBERED_HEADING_PATTERN.match(cleaned)
        if numbered_match is not None:
            cleaned = self._normalize_text(numbered_match.group("body"))
        cleaned = re.sub(r"^[（(]?[一二三四五六七八九十0-9]+[)）.、]\s*", "", cleaned)
        cleaned = cleaned.strip(" -、:：;；，。")
        return cleaned

    def _looks_like_requirement_topic(self, title: str, body: str, clause_code: str) -> bool:
        if not title:
            return False
        if self._is_generic_heading(title):
            return False
        if clause_code:
            return True
        if len(title) <= 4:
            return False
        if self._contains_topic_keyword(title):
            return True
        return any(token in body for token in ("优化", "改造", "升级", "开发", "建设", "实施"))

    def _contains_topic_keyword(self, text: str) -> bool:
        return any(keyword in text for keyword in self._SYSTEM_KEYWORDS)

    def _normalize_topic_key(self, title: str) -> str:
        normalized = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "", title)
        return normalized[:48] or "topic"

    def _expand_candidate_titles(self, raw_text: str, primary_title: str) -> list[str]:
        titles = [primary_title]
        for title in self._extract_mentioned_subsystems(raw_text):
            if self._normalize_topic_key(title) != self._normalize_topic_key(primary_title):
                titles.append(title)
        return titles[:6]

    def _extract_mentioned_subsystems(self, text: str) -> list[str]:
        body = self._prepare_topic_body(text)
        search_zone = body
        for marker in ("包括", "包含", "涵盖", "涉及", "由", "含"):
            if marker in search_zone:
                search_zone = search_zone.split(marker, 1)[1]
                break

        matches = self._SUBSYSTEM_PATTERN.findall(search_zone)
        results: list[str] = []
        for match in matches:
            title = self._normalize_subsystem_title(match)
            if not self._looks_like_subsystem_name(title):
                continue
            if title not in results:
                results.append(title)
        return results

    def _looks_like_subsystem_name(self, title: str) -> bool:
        if len(title) < 4:
            return False
        if title in {"系统", "控制系统", "子系统", "平台", "模块", "装置"}:
            return False
        return self._contains_topic_keyword(title)

    def _prepare_topic_body(self, text: str) -> str:
        body = self._strip_leading_codes(self._normalize_text(text))
        for prefix in self._NOISE_PREFIXES:
            if body.startswith(prefix):
                body = body[len(prefix) :]
                break
        body = self._normalize_subsystem_title(body)
        return body

    def _normalize_subsystem_title(self, text: str) -> str:
        body = self._strip_leading_codes(self._normalize_text(text))
        for prefix in self._CONTEXT_PREFIXES:
            if body.startswith(prefix):
                body = body[len(prefix) :]
                break
        trimmed = True
        while trimmed:
            trimmed = False
            for prefix in self._ACTION_PREFIXES:
                if body.startswith(prefix) and len(body) > len(prefix) + 3:
                    body = body[len(prefix) :].lstrip("的")
                    trimmed = True
                    break
        body = body.strip(" -、:：;；，。")
        return body

    def _strip_leading_codes(self, text: str) -> str:
        body = self._normalize_text(text)
        while True:
            matched_clause = self._CLAUSE_PATTERN.match(body)
            if matched_clause is None:
                break
            body = self._normalize_text(matched_clause.group("body"))
        return body

    def _is_generic_heading(self, text: str) -> bool:
        chapter_match = self._CHAPTER_PATTERN.match(self._normalize_text(text))
        if chapter_match is not None:
            title = self._normalize_text(chapter_match.group("title"))
            if not title:
                return True
            text = title
        return text in self._GENERIC_HEADINGS or len(text) <= 3

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()

    def _extract_standards(self, lines: list[str]) -> list[str]:
        pattern = re.compile(r"(?<![A-Za-z0-9])(?:GB|ISO|IEC|DL/T)[A-Za-z0-9\-/]*(?![A-Za-z0-9])", re.IGNORECASE)
        standards: list[str] = []
        for line in lines:
            found = pattern.findall(line)
            for item in found:
                normalized = item.upper()
                if normalized not in standards:
                    standards.append(normalized)
            if len(standards) >= 20:
                break
        return standards

    def _extract_acceptance(self, lines: list[str]) -> list[str]:
        acceptance = [line for line in lines if "验收" in line][:20]
        return acceptance
