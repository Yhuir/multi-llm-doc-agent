import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordExporterAgent:
    def __init__(self):
        self.name = "Layout & Word Export Agent"

    def run(self, final_data: dict, output_path: str = "outputs/工程实施方案.docx"):
        print(f"[{self.name}] 正在开始排版导出 Word (含 Mermaid 图表)...")
        doc = Document()
        
        # 1. 项目名称
        project_info = final_data.get("project_info", {})
        title = project_info.get("project_name", "工程实施方案")
        doc.add_heading(title, 0)
        
        # 2. 总体进度甘特图 (如果有渲染好的图片)
        if "master_gantt_img" in final_data:
            doc.add_heading("一、项目总体进度计划", level=1)
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(final_data["master_gantt_img"], width=Inches(6.0))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("图 1-1 项目总体实施进度甘特图").alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 3. 子系统循环
        for sub in final_data.get("subsystem_details", []):
            doc.add_heading(sub["name"], level=1)
            
            # 如果子系统有甘特图图片
            if sub.get("subsystem_gantt_img"):
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(sub["subsystem_gantt_img"], width=Inches(5.5))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f"图：{sub['name']} 细化进度图").alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for plan in sub.get("plans", []):
                doc.add_heading(plan["plan_title"], level=2)
                for content in plan.get("contents", []):
                    doc.add_heading(content["title"], level=3)
                    para = doc.add_paragraph(content["text"])
                    
                    # 插入章节配图 (如豆包生成的)
                    for img in content.get("images", []):
                        if img.get("source") and os.path.exists(img["source"]):
                            try:
                                doc.add_picture(img["source"], width=Inches(5.0))
                                doc.add_paragraph(img.get("caption", "示意图")).alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except: pass
        
        doc.save(output_path)
        print(f"[{self.name}] Word 文档导出成功: {output_path}")
        return output_path
