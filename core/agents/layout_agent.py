from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

class LayoutAgent:
    def __init__(self, output_dir="artifacts/final"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def _set_font_style(self, run, font_name="宋体", size_pt=12, bold=False, color=None):
        from docx.oxml.ns import qn
        run.font.name = font_name
        # For setting east asian font in docx
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size_pt)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    def generate_word(self, task_id: str, toc: dict, nodes_text: dict, images_meta_map: dict = None) -> str:
        """
        nodes_text: Dict mapping node_id to NodeText dictionary
        images_meta_map: Dict mapping node_id to list of image metadata dicts
        """
        from docx.oxml.ns import qn
        if images_meta_map is None:
            images_meta_map = {}
            
        doc = Document()
        
        # Keep track of global image numbering if needed, but we can do per-node or just sequentially.
        global_image_counter = 1
        
        def process_node(node, level):
            nonlocal global_image_counter
            # Heading
            if level == 1:
                h = doc.add_heading(node['title'], level=1)
                for run in h.runs:
                    self._set_font_style(run, '黑体', 16, True)
            elif level == 2:
                h = doc.add_heading(node['title'], level=2)
                for run in h.runs:
                    self._set_font_style(run, '黑体', 14, True)
            elif level == 3:
                h = doc.add_heading(node['title'], level=3)
                for run in h.runs:
                    self._set_font_style(run, '黑体', 12, True)
            elif level == 4:
                h = doc.add_heading(node['title'], level=4)
                for run in h.runs:
                    self._set_font_style(run, '黑体', 12, True)
                    
                node_id = node['node_id']
                node_content = nodes_text.get(node_id)
                node_images = images_meta_map.get(node_id, [])
                
                if node_content:
                    for section in node_content.get('sections', []):
                        p_h = doc.add_paragraph()
                        r_h = p_h.add_run(section.get('h', ''))
                        self._set_font_style(r_h, '黑体', 12, True)
                        
                        text = section.get('text', '')
                        if text:
                            p = doc.add_paragraph()
                            r = p.add_run(text)
                            self._set_font_style(r, '宋体', 12)
                            
                            # Handle "关键重难点"
                            if "重难点" in section.get('h', ''):
                                self._set_font_style(r, '宋体', 12, True, RGBColor(255, 0, 0))
                                
                        # Check if any image binds to this section (using section header 'h' or loosely matching)
                        # We will insert the image right after the section if bind_anchor is in the header or text
                        for img_meta in list(node_images): # copy list to allow removal
                            anchor = img_meta.get('bind_anchor', '')
                            if anchor and (anchor in section.get('h', '') or (text and anchor in text)):
                                img_path = img_meta.get('file')
                                if img_path and os.path.exists(img_path):
                                    # Insert Image
                                    doc.add_paragraph() # spacing
                                    p_img = doc.add_paragraph()
                                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    r_img = p_img.add_run()
                                    r_img.add_picture(img_path, width=Inches(5.0))
                                    
                                    # Insert Caption
                                    p_cap = doc.add_paragraph()
                                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    r_cap = p_cap.add_run(f"图{global_image_counter} {img_meta.get('caption', '')}")
                                    self._set_font_style(r_cap, '宋体', 10, False)
                                    global_image_counter += 1
                                    
                                    # Remove so we don't insert twice
                                    node_images.remove(img_meta)
                                    
                    # Insert any remaining images at the end of the node
                    for img_meta in node_images:
                        img_path = img_meta.get('file')
                        if img_path and os.path.exists(img_path):
                            doc.add_paragraph()
                            p_img = doc.add_paragraph()
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r_img = p_img.add_run()
                            r_img.add_picture(img_path, width=Inches(5.0))
                            
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r_cap = p_cap.add_run(f"图{global_image_counter} {img_meta.get('caption', '')}")
                            self._set_font_style(r_cap, '宋体', 10, False)
                            global_image_counter += 1

            for child in node.get('children', []):
                process_node(child, level + 1)
                
        for node in toc.get('tree', []):
            process_node(node, 1)
            
        output_path = os.path.join(self.output_dir, f"{task_id}_output.docx")
        doc.save(output_path)
        return output_path
