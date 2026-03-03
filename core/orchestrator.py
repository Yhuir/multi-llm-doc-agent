import uuid
import os
import json
import docx
import PyPDF2
from core.state_manager import StateManager
from utils.llm_client import LLMClient
from utils.doubao_client import DoubaoClient
from core.agents.parser_agent import RequirementParserAgent
from core.agents.toc_agent import TOCGeneratorAgent, TOCReviewAgent
from core.agents.writer_agent import SectionWriterAgent
from core.agents.layout_agent import LayoutAgent
from core.agents.length_agent import LengthControlAgent
from core.agents.image_agent import ImagePipelineAgent
from core.agents.style_extractor_agent import StyleExtractorAgent

class Orchestrator:
    def __init__(self, db_path="sqlite:///db/app.db", model_provider="deepseek-reasoner"):
        self.state_manager = StateManager(db_path)
        
        # 文本生成模型 (DeepSeek-V3.2 思考模式 或 豆包)
        self.llm_client = LLMClient(provider=model_provider) 
        
        # 图片生成模型 (豆包)
        self.doubao_client = DoubaoClient() 
        
        # 初始化各 Agent，注入对应的 Client
        self.parser_agent = RequirementParserAgent(self.llm_client)
        self.toc_generator = TOCGeneratorAgent(self.llm_client)
        self.toc_reviewer = TOCReviewAgent(self.llm_client)
        self.writer_agent = SectionWriterAgent(self.llm_client)
        self.layout_agent = LayoutAgent()
        self.length_agent = LengthControlAgent(self.llm_client)
        self.image_agent = ImagePipelineAgent(self.doubao_client)
        self.style_extractor = StyleExtractorAgent(self.llm_client)
        
        self.template1_text = ""
        self.template2_text = ""
        self.template2_path = "template/太和曲靖技术部分(1).pdf" # Default V3 template
        self._load_default_templates()

    def _load_default_templates(self):
        t1_path = "template/昆烟实施方案-目标范本.docx"
        t2_path = self.template2_path
        
        if os.path.exists(t1_path):
            try:
                doc = docx.Document(t1_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                self.template1_text = "\n".join(full_text)[:10000] # load up to 10k chars
            except Exception as e:
                print(f"Failed to load docx template: {e}")
                
        if os.path.exists(t2_path):
            try:
                reader = PyPDF2.PdfReader(t2_path)
                text = ""
                for page in reader.pages[:10]: # Read first 10 pages max
                    text += page.extract_text()
                self.template2_text = text[:10000]
            except Exception as e:
                print(f"Failed to load pdf template: {e}")

    def update_templates(self, t1_path: str, t2_path: str):
        if t1_path and os.path.exists(t1_path):
            try:
                if t1_path.endswith('.docx'):
                    doc = docx.Document(t1_path)
                    self.template1_text = "\n".join([p.text for p in doc.paragraphs])[:10000]
                elif t1_path.endswith('.pdf'):
                    reader = PyPDF2.PdfReader(t1_path)
                    self.template1_text = "".join([p.extract_text() for p in reader.pages[:10]])[:10000]
            except Exception as e:
                print(f"Failed to load t1: {e}")
                
        if t2_path and os.path.exists(t2_path):
            try:
                if t2_path.endswith('.docx'):
                    doc = docx.Document(t2_path)
                    self.template2_text = "\n".join([p.text for p in doc.paragraphs])[:10000]
                elif t2_path.endswith('.pdf'):
                    reader = PyPDF2.PdfReader(t2_path)
                    self.template2_text = "".join([p.extract_text() for p in reader.pages[:10]])[:10000]
            except Exception as e:
                print(f"Failed to load t2: {e}")

    def start_new_task(self, file_path: str) -> str:
        task_id = str(uuid.uuid4())
        self.state_manager.create_task(task_id, file_path)
        return task_id

    def process_parsing(self, task_id: str):
        task = self.state_manager.get_task(task_id)
        if not task:
            raise ValueError("Task not found")
            
        parsed_req = self.parser_agent.parse(task.file_path)
        
        os.makedirs(f"artifacts/{task_id}", exist_ok=True)
        with open(f"artifacts/{task_id}/requirement.json", "w", encoding='utf-8') as f:
            json.dump(parsed_req, f, ensure_ascii=False, indent=2)
            
        # V3: Style Extraction
        self.process_style_extraction(task_id)
            
        self.state_manager.update_task_status(task_id, "PARSED")
        return parsed_req

    def process_style_extraction(self, task_id: str):
        """Extract style from PDF template for V3."""
        style_dir = f"artifacts/{task_id}/style"
        os.makedirs(style_dir, exist_ok=True)
        
        # Use default template for now
        pdf_path = self.template2_path 
        if not os.path.exists(pdf_path):
            # Fallback to current dir if template folder is missing
            pdf_path = os.path.basename(pdf_path)
            
        if os.path.exists(pdf_path):
            style_profile = self.style_extractor.generate_style_profile(pdf_path)
            with open(f"{style_dir}/style_profile.json", "w", encoding='utf-8') as f:
                f.write(style_profile.model_dump_json(indent=2))
        else:
            # Save default style if no template found
            style_profile = self.style_extractor.get_default_style_profile()
            with open(f"{style_dir}/style_profile.json", "w", encoding='utf-8') as f:
                f.write(style_profile.model_dump_json(indent=2))

    def generate_toc(self, task_id: str):
        with open(f"artifacts/{task_id}/requirement.json", "r", encoding='utf-8') as f:
            parsed_req = json.load(f)
            
        toc = self.toc_generator.generate(parsed_req)
        self.state_manager.save_toc(task_id, toc.get('version', 1), toc)
        self.state_manager.update_task_status(task_id, "TOC_REVIEW")
        return toc

    def revise_toc(self, task_id: str, user_feedback: str):
        latest_toc_record = self.state_manager.get_latest_toc(task_id)
        if not latest_toc_record:
            raise ValueError("No TOC found to revise")
            
        new_toc = self.toc_reviewer.revise(latest_toc_record.toc_data, user_feedback)
        self.state_manager.save_toc(task_id, new_toc.get('version', latest_toc_record.version + 1), new_toc)
        return new_toc

    def confirm_toc_and_start_generation(self, task_id: str):
        self.state_manager.update_task_status(task_id, "GENERATING")
        latest_toc_record = self.state_manager.get_latest_toc(task_id)
        
        l3_nodes = []
        def extract_execution_nodes(node, level=1):
            if level == 4:
                l3_nodes.append(node)
            for child in node.get('children', []):
                extract_execution_nodes(child, level + 1)
                
        for node in latest_toc_record.toc_data.get('tree', []):
            extract_execution_nodes(node)
            
        for node in l3_nodes:
            self.state_manager.update_node_state(task_id, node['node_id'], "NODE_PENDING")
            
        return len(l3_nodes)

    def generate_content_for_node(self, task_id: str, node_id: str):
        self.state_manager.update_node_state(task_id, node_id, "TEXT_GENERATING")
        
        with open(f"artifacts/{task_id}/requirement.json", "r", encoding='utf-8') as f:
            parsed_req = json.load(f)
            
        latest_toc_record = self.state_manager.get_latest_toc(task_id)
        node_info = None
        def find_node(node):
            nonlocal node_info
            if node.get('node_id') == node_id:
                node_info = node
            for child in node.get('children', []):
                find_node(child)
        for node in latest_toc_record.toc_data.get('tree', []):
            find_node(node)
            
        if not node_info:
            self.state_manager.update_node_state(task_id, node_id, "NODE_FAILED", {"error": "Node not found in TOC"})
            return None
            
        # 1. Generate Base Text
        node_text = self.writer_agent.write_node(node_info, parsed_req, self.template1_text, self.template2_text)
        
        # 2. Length Control
        adjusted_text, word_count, length_pass = self.length_agent.adjust_length(node_text)
        self.state_manager.update_node_state(task_id, node_id, "LENGTH_CHECKED", {"word_count": word_count})
        
        os.makedirs(f"artifacts/{task_id}/nodes/{node_id}", exist_ok=True)
        with open(f"artifacts/{task_id}/nodes/{node_id}/text.json", "w", encoding='utf-8') as f:
            json.dump(adjusted_text, f, ensure_ascii=False, indent=2)

        # 3. Image Generation
        self.state_manager.update_node_state(task_id, node_id, "IMAGES_GENERATING")
        prompts_res = self.image_agent.generate_prompts(adjusted_text)
        images_meta = self.image_agent.generate_images_with_doubao(task_id, node_id, prompts_res.get('prompts', []))
        
        # Enrich images_meta with bind_anchor from adjusted_text.image_configs
        for img_meta in images_meta:
            img_id = img_meta['image_id']
            anchor = next((c.get('bind_anchor') for c in adjusted_text.get('image_configs', []) if c.get('image_id') == img_id), "未知锚点")
            img_meta['bind_anchor'] = anchor
            img_meta['caption'] = f"图 {img_id}"

        # 4. Image-Text Relevance Check
        relevance_res = self.image_agent.check_relevance(adjusted_text, images_meta)
        
        with open(f"artifacts/{task_id}/nodes/{node_id}/images.json", "w", encoding='utf-8') as f:
            json.dump(images_meta, f, ensure_ascii=False, indent=2)
            
        with open(f"artifacts/{task_id}/nodes/{node_id}/metrics.json", "w", encoding='utf-8') as f:
            json.dump({"word_count": word_count, "image_relevance": relevance_res}, f, ensure_ascii=False, indent=2)

        self.state_manager.update_node_state(task_id, node_id, "IMAGE_TEXT_CHECKED")
        self.state_manager.update_node_state(task_id, node_id, "TEXT_GENERATED")
        return adjusted_text

    def layout_document(self, task_id: str):
        self.state_manager.update_task_status(task_id, "LAYOUTING")
        latest_toc_record = self.state_manager.get_latest_toc(task_id)
        toc_data = latest_toc_record.toc_data
        
        # Load style profile
        style_profile = None
        style_path = f"artifacts/{task_id}/style/style_profile.json"
        if os.path.exists(style_path):
            with open(style_path, "r", encoding='utf-8') as f:
                style_profile = json.load(f)
        
        nodes_text = {}
        images_meta_map = {}
        
        nodes_dir = f"artifacts/{task_id}/nodes"
        if os.path.exists(nodes_dir):
            for node_id in os.listdir(nodes_dir):
                text_file = os.path.join(nodes_dir, node_id, "text.json")
                if os.path.exists(text_file):
                    with open(text_file, "r", encoding='utf-8') as f:
                        nodes_text[node_id] = json.load(f)
                        
                images_file = os.path.join(nodes_dir, node_id, "images.json")
                if os.path.exists(images_file):
                    with open(images_file, "r", encoding='utf-8') as f:
                        images_meta_map[node_id] = json.load(f)
                        
        template_path = self.template2_path.replace('.pdf', '.docx')
        if not os.path.exists(template_path):
            template_path = "template/standard_template.docx"
             
        output_path = self.layout_agent.generate_word(
            task_id, 
            toc_data, 
            nodes_text, 
            images_meta_map,
            template_path=template_path,
            style_profile=style_profile
        )
        self.state_manager.update_task_status(task_id, "DONE")
        return output_path
